from __future__ import annotations

import csv
import hashlib
import json
import math
import re
import shutil
import subprocess
import textwrap
import zipfile
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable


MISSING_TOKENS = {
    "",
    "-",
    "--",
    "na",
    "n/a",
    "nan",
    "none",
    "null",
    "missing",
    "unknown",
    "absent",
    "not collected",
    "not provided",
    "not applicable",
    "not available",
    "unavailable",
    "undetermined",
    "unidentified",
    "unspecified",
    "not reported",
    "not determined",
}

COUNTRY_STD_FIELDS = ("Country", "country")
COUNTRY_RAW_FIELDS = ("Geographic Location", "geographic_location", "geo_loc_name", "Country_Raw")
HOST_STD_FIELDS = ("Host_SD", "host_standardized", "Host Standardized")
HOST_RAW_FIELDS = ("Host", "host")
SOURCE_STD_FIELDS = ("Isolation_Source_SD", "isolation_source_standardized", "Isolation Source Standardized")
SOURCE_RAW_FIELDS = ("Isolation Source", "isolation_source")
SAMPLE_STD_FIELDS = ("Sample_Type_SD", "sample_type_standardized")
SAMPLE_RAW_FIELDS = ("Sample Type", "sample_type", "sample type", "Sample_Type")
ENV_STD_FIELDS = ("Environment_Medium_SD", "Environment_Broad_Scale_SD", "Environment_Local_Scale_SD")
ENV_RAW_FIELDS = ("Environment", "environment", "env_broad_scale", "env_local_scale", "env_medium")
BIOPROJECT_FIELDS = ("Assembly BioProject Accession", "BioProject Accession", "bioproject_accession")
ASSEMBLY_FIELDS = ("Assembly Accession", "assembly_accession")
ORGANISM_FIELDS = ("Organism Name", "organism_name")
ASSEMBLY_LEVEL_FIELDS = ("Assembly Level", "assembly_level")
RELEASE_DATE_FIELDS = ("Assembly Release Date", "release_date")
COLLECTION_DATE_FIELDS = ("Collection Date", "collection_date")
COMPLETENESS_FIELDS = ("CheckM completeness", "Completeness", "CheckM2_Completeness", "completeness")
CONTAMINATION_FIELDS = ("CheckM contamination", "Contamination", "CheckM2_Contamination", "contamination")
N50_FIELDS = ("Assembly Stats Contig N50", "Assembly Stats Scaffold N50", "N50")
CONTIG_FIELDS = ("Assembly Stats Number of Contigs", "contigs", "Contigs")
GENOME_SIZE_FIELDS = ("Assembly Stats Total Sequence Length", "Assembly Stats Total Ungapped Length", "Genome size")
GC_FIELDS = ("Assembly Stats GC Percent", "GC", "GC%")
COUNTRY_CONFIDENCE_FIELDS = ("Country_Confidence", "Country_Evidence")
HOST_CONFIDENCE_FIELDS = ("Host_SD_Confidence", "Host_SD_Method")

POLITICAL_CONTINENTS = {
    "Africa",
    "Antarctica",
    "Asia",
    "Europe",
    "North America",
    "Oceania",
    "South America",
}

POLITICAL_SUBCONTINENTS = {
    "Australia and New Zealand",
    "Caribbean",
    "Central America",
    "Central Asia",
    "Eastern Africa",
    "Eastern Asia",
    "Eastern Europe",
    "Melanesia",
    "Micronesia",
    "Middle Africa",
    "Northern Africa",
    "Northern America",
    "Northern Asia",
    "Northern Europe",
    "Polynesia",
    "South America",
    "South-Eastern Asia",
    "Southern Africa",
    "Southern Asia",
    "Southern Europe",
    "Western Africa",
    "Western Asia",
    "Western Europe",
    "Antarctica",
}


@dataclass
class TaxonInput:
    id: int
    name: str
    rank: str
    genome_count: int | None
    metadata_clean_path: str
    last_synced_at: str | None = None


@dataclass
class TaxonStats:
    name: str
    rank: str
    rows: int = 0
    country_usable: int = 0
    host_or_source_usable: int = 0
    year_usable: int = 0
    quality_available: int = 0
    isolation_usable: int = 0
    confidence_available: int = 0
    bioprojects: Counter[str] = field(default_factory=Counter)
    countries: Counter[str] = field(default_factory=Counter)
    hosts: Counter[str] = field(default_factory=Counter)
    sources: Counter[str] = field(default_factory=Counter)
    years: Counter[str] = field(default_factory=Counter)


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def safe_slug(value: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9._-]+", "-", value.strip()).strip("-").lower()
    return slug or "global-insights"


def normalize(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip())


def normalized_key(value: Any) -> str:
    return normalize(value).casefold()


def is_usable(value: Any) -> bool:
    text = normalized_key(value)
    if text in MISSING_TOKENS:
        return False
    if not text:
        return False
    if text.startswith("unknown ") or text.endswith(" unknown"):
        return False
    if text.startswith("absent ") or text.endswith(" absent"):
        return False
    return True


def safe_display_value(value: Any, fallback: str = "missing/unusable") -> str:
    text = normalize(value)
    return text if is_usable(text) else fallback


def evidence_value(row: dict[str, Any], fields: Iterable[str]) -> str:
    for field_name in fields:
        value = row.get(field_name)
        if is_usable(value):
            return normalize(value)
    return "not recorded"


def correction_evidence(row: dict[str, Any], field_name: str, raw_value: str, standardized_value: str) -> tuple[str, str, str]:
    if field_name == "Country":
        evidence = evidence_value(row, ("Country_Source", "Country_Evidence", "Country_Confidence"))
        confidence = evidence_value(row, ("Country_Confidence", "Country_Evidence"))
    elif field_name == "Host":
        evidence = evidence_value(row, ("Host_SD_Method", "Host_SD_Source", "Host_Source", "Host_SD_Evidence"))
        confidence = evidence_value(row, ("Host_SD_Confidence", "Host_Confidence", "Host_SD_Method"))
    elif field_name == "Isolation source":
        evidence = evidence_value(row, ("Isolation_Source_SD_Method", "Isolation_Source_Source", "Isolation_Source_Evidence"))
        confidence = evidence_value(row, ("Isolation_Source_SD_Confidence", "Isolation_Source_Confidence"))
    elif field_name == "Sample type":
        evidence = evidence_value(row, ("Sample_Type_SD_Method", "Sample_Type_Source", "Sample_Type_Evidence"))
        confidence = evidence_value(row, ("Sample_Type_SD_Confidence", "Sample_Type_Confidence"))
    elif field_name == "Environment":
        evidence = evidence_value(row, ("Environment_Medium_SD_Method", "Environment_Source", "Environment_Evidence"))
        confidence = evidence_value(row, ("Environment_Medium_SD_Confidence", "Environment_Confidence"))
    else:
        evidence = "not recorded"
        confidence = "not recorded"

    if not is_usable(raw_value) and is_usable(standardized_value):
        raw_display = f"primary {field_name.lower()} field absent/unusable; recovered from secondary evidence"
        if evidence == "not recorded":
            evidence = "secondary standardized metadata field"
    else:
        raw_display = safe_display_value(raw_value)
    return raw_display[:180], evidence[:180], confidence[:120]


def correction_type(field_name: str, raw_value: str, standardized_value: str, evidence: str = "") -> str:
    raw_ok = is_usable(raw_value)
    std_ok = is_usable(standardized_value)
    evidence_key = normalized_key(evidence)
    if not std_ok:
        return "unresolved_or_not_standardized"
    if not raw_ok:
        return "secondary_evidence_recovery"
    if "block" in evidence_key or "non-source" in evidence_key:
        return "blocked_or_non_source_descriptor"
    if field_name in {"Environment", "Sample type"}:
        return "derived_context_assignment"
    compact_raw = re.sub(r"[^a-z0-9]+", "", normalized_key(raw_value))
    compact_std = re.sub(r"[^a-z0-9]+", "", normalized_key(standardized_value))
    if compact_raw == compact_std and normalized_key(raw_value) != normalized_key(standardized_value):
        return "spelling_or_format_normalization"
    if field_name in {"Country", "Host"}:
        return "synonym_normalization"
    return "controlled_category_mapping"


def confidence_bucket(raw_value: str, standardized_value: str, confidence: str = "", evidence: str = "") -> str:
    if not is_usable(standardized_value):
        return "unresolved"
    if not is_usable(raw_value):
        return "secondary-evidence-recovered"
    text = normalized_key(f"{confidence} {evidence}")
    if any(token in text for token in ("needs review", "review-needed", "ambiguous")):
        return "review-needed"
    if any(token in text for token in ("manual", "approved", "exact", "high", "taxonomy")):
        return "high"
    if any(token in text for token in ("synonym", "dictionary", "rule", "medium", "controlled")):
        return "medium"
    if any(token in text for token in ("fuzzy", "low", "suggest", "candidate")):
        return "low"
    return "confidence-not-recorded"


def readiness_tier(row: dict[str, Any]) -> tuple[str, str]:
    assemblies = int(row.get("assemblies") or 0)
    country = float(row.get("country_completeness_percent") or 0)
    host_source = float(row.get("host_or_source_completeness_percent") or 0)
    year = float(row.get("collection_year_completeness_percent") or 0)
    if assemblies >= 1000 and country >= 80 and host_source >= 80 and year >= 70:
        return "strict", ">=1,000 assemblies, >=80% country, >=80% host/source, >=70% collection year"
    if assemblies >= 100 and country >= 70 and host_source >= 70 and year >= 50:
        return "standard", ">=100 assemblies, >=70% country, >=70% host/source, >=50% collection year"
    if assemblies >= 50 and country >= 50 and host_source >= 50:
        return "exploratory", ">=50 assemblies, >=50% country and host/source"
    return "not-ready", "Below exploratory metadata-readiness thresholds"


def denominator_note(total: int | float, scope: str = "all non-redundant assemblies") -> str:
    return f"Denominator: {int(total or 0):,} {scope}."


def row_value(row: dict[str, Any], fields: Iterable[str]) -> str:
    for field_name in fields:
        value = row.get(field_name)
        if value is not None and str(value).strip():
            return normalize(value)
    return ""


def parse_year(value: str) -> str:
    text = normalize(value)
    if not text:
        return ""
    match = re.search(r"(19|20)\d{2}", text)
    if not match:
        return ""
    year = int(match.group(0))
    if 1900 <= year <= datetime.now(timezone.utc).year + 1:
        return str(year)
    return ""


NON_CANONICAL_SPECIES_TOKENS = {
    "sp",
    "sp.",
    "spp",
    "spp.",
    "bacterium",
    "archaeon",
    "microorganism",
    "metagenome",
    "uncultured",
    "unclassified",
    "endosymbiont",
    "symbiont",
}

TAXONOMY_LABEL_CLASS_ORDER = (
    "canonical_species",
    "unresolved_species_level_label",
    "provisional_taxonomic_label",
    "noncanonical_species_label",
    "canonical_genus",
    "unknown_taxon_label",
)

TAXONOMY_LABEL_CLASS_LABELS = {
    "canonical_species": "Canonical species",
    "unresolved_species_level_label": "Unresolved species-level label",
    "provisional_taxonomic_label": "Provisional taxonomic label",
    "noncanonical_species_label": "Non-canonical species label",
    "canonical_genus": "Canonical genus",
    "unknown_taxon_label": "Unknown taxon label",
}

TAXONOMY_LABEL_CLASS_DESCRIPTIONS = {
    "canonical_species": "Binomial species label used for standard species-level reporting.",
    "unresolved_species_level_label": "Placeholder, strain-like, uncultured, metagenomic, or unclassified species-level label; searchable but not counted as a canonical species.",
    "provisional_taxonomic_label": "Candidatus or provisional nomenclature; searchable but reported separately from canonical species.",
    "noncanonical_species_label": "Species-level label that does not meet canonical binomial rules; searchable but reported separately.",
    "canonical_genus": "Genus-level taxon page.",
    "unknown_taxon_label": "Taxon label is missing or unusable.",
}


def taxonomy_label_metadata(name: str, rank: str = "species") -> dict[str, str]:
    normalized = normalize(name)
    rank_key = normalize(rank).casefold()
    if rank_key == "genus":
        key = "canonical_genus"
    elif not normalized:
        key = "unknown_taxon_label"
    else:
        lower = normalized.casefold()
        parts = normalized.split()
        if parts and parts[0].casefold() == "candidatus":
            key = "provisional_taxonomic_label"
        elif "uncultured" in lower or "metagenome" in lower or "unclassified" in lower:
            key = "unresolved_species_level_label"
        elif any(token in lower for token in (" sp.", " spp.", " species complex", " group ", " clade ")) or lower.endswith(" sp") or re.search(r"\bsp\d", lower):
            key = "unresolved_species_level_label"
        elif any(char in normalized for char in "()[]"):
            key = "noncanonical_species_label"
        elif len(parts) >= 2:
            epithet = parts[1].strip().rstrip(".,;:")
            if re.match(r"^[a-z][a-z0-9-]*$", epithet.casefold()) and epithet.casefold() not in NON_CANONICAL_SPECIES_TOKENS and not re.match(r"^sp\d", epithet.casefold()):
                key = "canonical_species"
            else:
                key = "noncanonical_species_label"
        else:
            key = "noncanonical_species_label"
    return {
        "key": key,
        "label": TAXONOMY_LABEL_CLASS_LABELS[key],
        "description": TAXONOMY_LABEL_CLASS_DESCRIPTIONS[key],
    }


def parse_taxonomy(organism_name: str, fallback_taxon: str = "") -> tuple[str, str]:
    text = normalize(organism_name) or normalize(fallback_taxon)
    if not text:
        return "Unclassified", "Unclassified"
    parts = [part for part in re.split(r"\s+", text) if part]
    if parts and parts[0].casefold() == "candidatus" and len(parts) >= 3:
        genus = f"{parts[0]} {parts[1]}"
        species = f"{parts[0]} {parts[1]} {parts[2]}"
    else:
        genus = parts[0] if parts else "Unclassified"
        if len(parts) >= 2:
            species = f"{parts[0]} {parts[1]}"
        else:
            species = genus
    return genus, species


def taxonomy_label_summary_rows(label_counts: Counter[str], label_name_sets: dict[str, set[str]], total_labels: int, total_assemblies: int | None = None) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for key in TAXONOMY_LABEL_CLASS_ORDER:
        label_count = len(label_name_sets.get(key, set()))
        assembly_count = int(label_counts.get(key, 0))
        if label_count <= 0 and assembly_count <= 0:
            continue
        row = {
            "label_class": key,
            "label": TAXONOMY_LABEL_CLASS_LABELS[key],
            "description": TAXONOMY_LABEL_CLASS_DESCRIPTIONS[key],
            "labels": label_count,
            "label_percent": percent(label_count, total_labels),
            "denominator": int(total_labels),
            "denominator_note": denominator_note(total_labels, "distinct species-level labels"),
        }
        if total_assemblies is not None:
            row.update(
                {
                    "assemblies": assembly_count,
                    "assembly_percent": percent(assembly_count, total_assemblies),
                    "assembly_denominator": int(total_assemblies),
                    "assembly_denominator_note": denominator_note(total_assemblies),
                }
            )
        rows.append(row)
    return rows


def parse_float(value: str) -> float | None:
    text = normalize(value).replace(",", "")
    if not text:
        return None
    try:
        number = float(text)
    except ValueError:
        return None
    if math.isnan(number) or math.isinf(number):
        return None
    return number


def percent(numerator: int | float, denominator: int | float) -> float:
    if denominator <= 0:
        return 0.0
    return round((float(numerator) / float(denominator)) * 100.0, 2)


def top_rows(counter: Counter[str], total: int, limit: int = 20) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for rank, (value, count) in enumerate(counter.most_common(limit), start=1):
        rows.append(
            {
                "rank": rank,
                "label": value,
                "count": int(count),
                "percent": percent(count, total),
                "denominator": int(total),
                "denominator_note": denominator_note(total),
            }
        )
    return rows


def severity_from_share(value: float) -> str:
    if value >= 75:
        return "severe"
    if value >= 50:
        return "high"
    if value >= 25:
        return "moderate"
    return "low"


def classify_host_source(host: str, source: str) -> str:
    host_text = normalize(host) if is_usable(host) else ""
    source_text = normalize(source) if is_usable(source) else ""
    text = f"{host_text} {source_text}".casefold()
    if not text.strip():
        return "Missing/ambiguous"
    if any(token in text for token in ("human", "homo sapiens", "patient", "clinical", "hospital", "stool", "feces", "faeces", "blood", "urine", "sputum")):
        return "Human-associated"
    if any(token in text for token in ("chicken", "cattle", "cow", "pig", "swine", "bird", "fish", "mouse", "rat", "dog", "cat", "animal", "poultry")):
        return "Animal-associated"
    if any(token in text for token in ("plant", "root", "leaf", "rhizosphere", "phyllosphere", "crop")):
        return "Plant-associated"
    if any(token in text for token in ("food", "meat", "milk", "cheese", "dairy", "vegetable", "produce", "seafood")):
        return "Food-associated"
    if "soil" in text:
        return "Soil"
    if any(token in text for token in ("water", "river", "lake", "marine", "seawater", "wastewater", "sewage", "aquatic")):
        return "Aquatic/water"
    if any(token in text for token in ("environment", "sediment", "biofilm", "surface", "air", "built environment")):
        return "Environmental"
    if any(token in text for token in ("lab", "laboratory", "culture medium")):
        return "Laboratory/culture"
    return "Other/ambiguous"


def resolve_metadata_path(path_text: str) -> Path:
    path = Path(path_text)
    if path.exists():
        return path
    marker = "/app/fetchm_webapp"
    if path_text.startswith(marker):
        local_root = Path(__file__).resolve().parents[1]
        candidate = local_root / path_text[len(marker) + 1 :]
        if candidate.exists():
            return candidate
    return path


def sorted_taxa(taxa: Iterable[dict[str, Any] | TaxonInput]) -> list[TaxonInput]:
    normalized: list[TaxonInput] = []
    for item in taxa:
        if isinstance(item, TaxonInput):
            normalized.append(item)
        else:
            normalized.append(
                TaxonInput(
                    id=int(item.get("id") or 0),
                    name=str(item.get("name") or item.get("species_name") or ""),
                    rank=str(item.get("rank") or item.get("taxon_rank") or "species"),
                    genome_count=int(item["genome_count"]) if item.get("genome_count") not in (None, "") else None,
                    metadata_clean_path=str(item.get("metadata_clean_path") or ""),
                    last_synced_at=str(item.get("last_synced_at") or "") or None,
                )
            )

    def sort_key(taxon: TaxonInput) -> tuple[int, str, str]:
        rank_order = 0 if taxon.rank == "species" else 1 if taxon.rank == "genus" else 2
        newest_first = "".join(chr(255 - ord(char)) for char in (taxon.last_synced_at or ""))
        return (rank_order, newest_first, taxon.name.casefold())

    return sorted(normalized, key=sort_key)


def quality_available(row: dict[str, Any]) -> bool:
    return any(is_usable(row_value(row, fields)) for fields in (COMPLETENESS_FIELDS, CONTAMINATION_FIELDS, N50_FIELDS, CONTIG_FIELDS, GENOME_SIZE_FIELDS, GC_FIELDS))


def confidence_available(row: dict[str, Any]) -> bool:
    return any(is_usable(row_value(row, fields)) for fields in (COUNTRY_CONFIDENCE_FIELDS, HOST_CONFIDENCE_FIELDS))


def update_field_pair_stats(stats: dict[str, int], raw_value: str, standardized_value: str) -> None:
    raw_ok = is_usable(raw_value)
    standardized_ok = is_usable(standardized_value)
    if raw_ok:
        stats["raw_usable"] += 1
    if standardized_ok:
        stats["standardized_usable"] += 1
    if raw_ok and standardized_ok:
        stats["both_usable"] += 1
        if normalized_key(raw_value) != normalized_key(standardized_value):
            stats["changed_mappings"] += 1
    elif standardized_ok:
        stats["standardized_only"] += 1
    elif raw_ok:
        stats["raw_only"] += 1


def write_csv(path: Path, rows: list[dict[str, Any]], fieldnames: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)




def load_validation_accuracy(path: Path | None = None) -> tuple[list[dict[str, Any]], bool]:
    """Summarize optional manually reviewed validation records when available."""
    source = path or Path(__file__).with_name("validation_records.csv")
    fields = ["Country", "Host", "Isolation source", "Sample type", "Environment", "Collection year from collection-date metadata"]
    if not source.exists():
        return [
            {
                "field": field,
                "validation_records": 0,
                "precision_percent": "not available",
                "false_positive_rate_percent": "not available",
                "unresolved_rate_percent": "not available",
                "common_error_types": "not available",
                "validation_source": str(source),
                "status": "manual validation CSV not available",
            }
            for field in fields
        ], False

    counters: dict[str, Counter[str]] = defaultdict(Counter)
    errors: dict[str, Counter[str]] = defaultdict(Counter)
    with source.open(newline="", encoding="utf-8", errors="replace") as handle:
        for row in csv.DictReader(handle):
            field = normalize(row.get("field")) or "Unknown"
            decision = normalized_key(row.get("reviewer_decision") or row.get("decision") or row.get("reviewer_label"))
            error_type = normalize(row.get("error_type"))
            counters[field]["total"] += 1
            if decision in {"correct", "accept", "accepted", "true_positive", "tp", "pass"}:
                counters[field]["correct"] += 1
            elif decision in {"false_positive", "incorrect", "wrong", "reject", "rejected", "fp", "fail"}:
                counters[field]["false_positive"] += 1
            elif decision in {"unresolved", "ambiguous", "unknown", "review"}:
                counters[field]["unresolved"] += 1
            else:
                counters[field]["unclassified"] += 1
            if error_type:
                errors[field][error_type] += 1

    rows: list[dict[str, Any]] = []
    for field, counter in sorted(counters.items()):
        total = int(counter.get("total", 0))
        rows.append(
            {
                "field": field,
                "validation_records": total,
                "precision_percent": percent(counter.get("correct", 0), max(counter.get("correct", 0) + counter.get("false_positive", 0), 1)),
                "false_positive_rate_percent": percent(counter.get("false_positive", 0), total),
                "unresolved_rate_percent": percent(counter.get("unresolved", 0), total),
                "common_error_types": "; ".join(f"{label}:{count}" for label, count in errors[field].most_common(5)) or "none recorded",
                "validation_source": str(source),
                "status": "available",
            }
        )
    return rows, bool(rows)


def build_field_confidence_summary(counters: dict[str, Counter[str]], denominator: int) -> list[dict[str, Any]]:
    statuses = ["high", "medium", "low", "review-needed", "unresolved", "secondary-evidence-recovered", "confidence-not-recorded"]
    rows: list[dict[str, Any]] = []
    for field in ["Country", "Host", "Isolation source", "Sample type", "Environment", "Collection year from collection-date metadata"]:
        counter = counters.get(field, Counter())
        total = sum(counter.values()) or denominator
        for status in statuses:
            count = int(counter.get(status, 0))
            rows.append(
                {
                    "field": field,
                    "confidence_status": status,
                    "count": count,
                    "percent": percent(count, total),
                    "denominator": int(total),
                    "denominator_note": denominator_note(total, f"records evaluated for {field}"),
                }
            )
    return rows


def build_case_studies(taxon_stats: dict[str, TaxonStats], taxon_quality: list[dict[str, Any]]) -> list[dict[str, Any]]:
    quality_by_name = {normalized_key(row.get("taxon")): row for row in taxon_quality}
    stats_by_name = {normalized_key(stat.name): stat for stat in taxon_stats.values()}
    selected = [
        "Salmonella enterica",
        "Escherichia coli",
        "Klebsiella pneumoniae",
        "Staphylococcus aureus",
        "Campylobacter jejuni",
        "Pseudomonas fluorescens",
    ]
    rows: list[dict[str, Any]] = []
    for name in selected:
        stat = stats_by_name.get(normalized_key(name))
        quality = quality_by_name.get(normalized_key(name), {})
        if not stat:
            rows.append({"taxon": name, "status": "not available in current snapshot", "assemblies": 0})
            continue
        top_project, top_project_count = stat.bioprojects.most_common(1)[0] if stat.bioprojects else ("not recorded", 0)
        rows.append(
            {
                "taxon": stat.name,
                "rank": stat.rank,
                "status": "available",
                "assemblies": stat.rows,
                "metadata_quality_score": quality.get("metadata_quality_score", ""),
                "metadata_quality_grade": quality.get("metadata_quality_grade", ""),
                "metadata_readiness_tier": quality.get("metadata_readiness_tier", ""),
                "country_completeness_percent": quality.get("country_completeness_percent", ""),
                "host_or_source_completeness_percent": quality.get("host_or_source_completeness_percent", ""),
                "collection_year_completeness_percent": quality.get("collection_year_completeness_percent", ""),
                "top_countries": "; ".join(f"{label}:{count}" for label, count in stat.countries.most_common(5)),
                "top_hosts": "; ".join(f"{label}:{count}" for label, count in stat.hosts.most_common(5)),
                "top_sources": "; ".join(f"{label}:{count}" for label, count in stat.sources.most_common(5)),
                "top_bioproject": top_project,
                "top_bioproject_share_percent": percent(top_project_count, stat.rows),
                "sampling_caution": "Consider BioProject-aware or country/host-balanced subsampling before comparative analysis." if percent(top_project_count, stat.rows) >= 25 else "No severe single-project dominance detected by the default threshold.",
                "denominator": stat.rows,
                "denominator_note": denominator_note(stat.rows, f"assemblies assigned to {stat.name}"),
            }
        )
    return rows


def write_json(path: Path, payload: dict[str, Any] | list[Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def write_checksums(base_dir: Path) -> str:
    rows: list[str] = []
    for path in sorted(base_dir.rglob("*")):
        if not path.is_file() or path.name == "checksums.sha256":
            continue
        relative = path.relative_to(base_dir).as_posix()
        rows.append(f"{sha256_file(path)}  {relative}")
    checksum_path = base_dir / "checksums.sha256"
    checksum_path.write_text("\n".join(rows) + ("\n" if rows else ""), encoding="utf-8")
    return "checksums.sha256"




def write_publication_export_audit(summary: dict[str, Any], snapshot_dir: Path) -> dict[str, Any]:
    checks: list[dict[str, Any]] = []
    for figure in summary.get("figure_exports") or []:
        figure_id = figure.get("figure_id", "unknown")
        files = figure.get("files") or {}
        expected = {
            "svg": files.get("svg"),
            "pdf": files.get("pdf"),
            "png": files.get("png"),
            "source_csv": figure.get("source_data"),
            "legend_txt": figure.get("legend_file"),
        }
        for kind, relative in expected.items():
            file_path = snapshot_dir / str(relative or "")
            checks.append(
                {
                    "figure_id": figure_id,
                    "artifact": kind,
                    "path": relative or "",
                    "exists": bool(relative) and file_path.exists() and file_path.is_file(),
                    "nonzero_size": bool(relative) and file_path.exists() and file_path.stat().st_size > 0,
                    "denominator_recorded": bool(figure.get("denominator")),
                    "png_expected_dpi": 600 if kind == "png" else "not applicable",
                }
            )
    passed = all(row["exists"] and row["nonzero_size"] and row["denominator_recorded"] for row in checks) and len(checks) >= 45
    audit = {
        "status": "pass" if passed else "warning",
        "snapshot_id": summary.get("snapshot_id"),
        "generated_at": utc_now(),
        "figure_count": len(summary.get("figure_exports") or []),
        "checks": checks,
    }
    write_json(snapshot_dir / "provenance" / "publication_export_audit.json", audit)
    return audit

def write_snapshot_manifest(summary: dict[str, Any], snapshot_dir: Path) -> None:
    methods = summary.get("methods") or {}
    provenance_dir = snapshot_dir / "provenance"
    manifest = {
        "snapshot_id": summary.get("snapshot_id"),
        "generated_at": summary.get("generated_at"),
        "summary_json": "summary.json",
        "reports_dir": "reports",
        "figures_dir": "figures",
        "source_data_dir": "source_data",
        "tables_dir": "tables",
        "provenance_dir": "provenance",
        "unique_assemblies": (summary.get("overview") or {}).get("unique_assemblies"),
        "metadata_rows_scanned": methods.get("metadata_rows_scanned"),
        "duplicate_rows_skipped": methods.get("duplicate_rows_skipped"),
        "app_version": methods.get("app_version"),
        "app_commit": methods.get("app_commit"),
        "standardization_rule_version": methods.get("standardization_rule_version"),
        "interpretation_caution": methods.get("caution"),
    }
    write_json(snapshot_dir / "manifest.json", manifest)
    write_json(provenance_dir / "manifest.json", manifest)
    write_json(provenance_dir / "software_versions.json", {
        "app_version": methods.get("app_version"),
        "app_commit": methods.get("app_commit"),
        "ncbi_datasets_version": methods.get("ncbi_datasets_version"),
        "taxonkit_version": methods.get("taxonkit_version"),
    })
    write_json(provenance_dir / "rule_fingerprint.json", {
        "standardization_rule_version": methods.get("standardization_rule_version"),
        "standardization_rule_rows": methods.get("standardization_rule_rows"),
        "standardization_approved_rule_rows": methods.get("standardization_approved_rule_rows"),
        "standardization_rule_files": methods.get("standardization_rule_files"),
        "latest_standardization_audit_timestamp": methods.get("latest_standardization_audit_timestamp"),
        "latest_standardization_audit_git_commit": methods.get("latest_standardization_audit_git_commit"),
        "latest_standardization_audit_code_version": methods.get("latest_standardization_audit_code_version"),
        "host_standardization_provenance": methods.get("host_standardization_provenance") or {},
        "geography_collection_date_provenance": methods.get("geography_collection_date_provenance") or {},
    })
    write_json(provenance_dir / "qa_report.json", summary.get("qa") or {})


def command_version(command: list[str]) -> str:
    try:
        result = subprocess.run(command, check=False, capture_output=True, text=True, timeout=8)
    except Exception:
        return "not available"
    output = normalize((result.stdout or result.stderr or "").splitlines()[0] if (result.stdout or result.stderr) else "")
    return output or f"available, return code {result.returncode}"


def standardization_root() -> Path:
    return Path(__file__).resolve().parents[1] / "standardization"


def latest_host_standardization_provenance(root: Path | None = None) -> dict[str, Any]:
    root = root or Path(__file__).resolve().parents[1]
    monitoring_root = root / "data" / "host_standardization_monitoring"
    latest_path = monitoring_root / "latest.json"
    if not latest_path.exists():
        return {}
    try:
        latest = json.loads(latest_path.read_text(encoding="utf-8"))
        summary_path = monitoring_root / str(latest.get("summary") or "")
        summary = json.loads(summary_path.read_text(encoding="utf-8"))
    except (OSError, ValueError, TypeError):
        return {}
    row_counts = summary.get("rule_row_counts") or {}
    return {
        "host_rule_version": summary.get("host_rule_version") or "not available",
        "host_rule_commit": summary.get("host_rule_commit") or "b92a591",
        "host_synonyms_row_count": int(row_counts.get("host_synonyms") or 0),
        "host_negative_rules_row_count": int(row_counts.get("host_negative_rules") or 0),
        "host_context_rules_row_count": int(row_counts.get("host_context_rules") or 0),
        "host_microbial_allowlist_row_count": int(row_counts.get("host_microbial_allowlist") or 0),
        "host_sd_microbial_leakage_count": int(summary.get("microbial_leakage_count") or 0),
        "latest_host_qa_timestamp": summary.get("generated_at") or "not available",
        "validation_sample_filename": summary.get("validation_sample_filename") or "",
        "validation_sample_sha256": summary.get("validation_sample_sha256") or "",
    }


def latest_geography_collection_date_provenance(root: Path | None = None) -> dict[str, Any]:
    root = root or Path(__file__).resolve().parents[1]
    qa_root = root / "data" / "geography_collection_date_qa"
    latest_path = qa_root / "latest.json"
    if not latest_path.exists():
        return {}
    try:
        latest = json.loads(latest_path.read_text(encoding="utf-8"))
        summary_path = qa_root / str(latest.get("summary") or "")
        summary = json.loads(summary_path.read_text(encoding="utf-8"))
    except (OSError, ValueError, TypeError):
        return {}
    metrics = summary.get("metrics") or {}
    provenance = summary.get("provenance") or {}
    return {
        "qa_timestamp": provenance.get("qa_timestamp") or "not available",
        "qa_commit": provenance.get("qa_commit") or "not available",
        "total_canonical_rows_audited": int(metrics.get("total_rows_scanned") or 0),
        "country_coverage_percent": float(metrics.get("country_present_percent") or 0),
        "continent_coverage_percent": float(metrics.get("continent_present_percent") or 0),
        "subcontinent_coverage_percent": float(metrics.get("subcontinent_present_percent") or 0),
        "collection_year_coverage_percent": float(metrics.get("collection_year_present_percent") or 0),
        "non_country_values_in_country_rows": int(metrics.get("non_country_values_in_country_rows") or 0),
        "country_continent_mismatch_rows": int(metrics.get("country_continent_mismatch_rows") or 0),
        "country_subcontinent_mismatch_rows": int(metrics.get("country_subcontinent_mismatch_rows") or 0),
        "invalid_future_impossible_collection_year_rows": (
            int(metrics.get("invalid_collection_year_rows") or 0)
            + int(metrics.get("future_collection_year_rows") or 0)
            + int(metrics.get("impossible_collection_year_rows") or 0)
        ),
        "country_lookup_sha256": provenance.get("country_lookup_sha256") or "",
        "collection_date_parser_sha256": provenance.get("collection_date_parser_sha256") or "",
        "generated_artifacts": provenance.get("generated_artifacts") or [],
    }


def latest_production_readiness_gate() -> dict[str, Any] | None:
    root = standardization_root() / "review" / "final_audit"
    if not root.exists():
        return None
    candidates = sorted(root.glob("*/production_readiness_gate.json"), key=lambda path: path.parent.name, reverse=True)
    for candidate in candidates:
        try:
            return json.loads(candidate.read_text(encoding="utf-8"))
        except Exception:
            continue
    return None


def standardization_rule_manifest() -> dict[str, Any]:
    root = standardization_root()
    rule_files = [
        root / "host_synonyms.csv",
        root / "host_negative_rules.csv",
        root / "host_context_rules.csv",
        root / "host_microbial_allowlist.csv",
        root / "controlled_categories.csv",
        root / "approved_broad_categories.csv",
        root / "geography_reviewed_rules.csv",
        root / "collection_date_reviewed_rules.csv",
    ]
    digest = hashlib.sha256()
    files: list[dict[str, Any]] = []
    total_rows = 0
    approved_rows = 0
    review_rows = 0
    for rule_file in rule_files:
        if not rule_file.exists():
            files.append({"path": str(rule_file.relative_to(root.parent)), "exists": False, "rows": 0})
            continue
        data = rule_file.read_bytes()
        digest.update(rule_file.name.encode("utf-8"))
        digest.update(data)
        rows = 0
        file_approved = 0
        file_review = 0
        try:
            with rule_file.open(newline="", encoding="utf-8", errors="replace") as handle:
                for row in csv.DictReader(handle):
                    rows += 1
                    status = normalized_key(row.get("status"))
                    if status == "approved":
                        file_approved += 1
                    elif status and status != "approved":
                        file_review += 1
        except Exception:
            rows = max(0, data.count(b"\n") - 1)
        total_rows += rows
        approved_rows += file_approved
        review_rows += file_review
        files.append(
            {
                "path": str(rule_file.relative_to(root.parent)),
                "exists": True,
                "rows": rows,
                "approved_rows": file_approved,
                "review_or_nonapproved_rows": file_review,
                "sha256_12": hashlib.sha256(data).hexdigest()[:12],
            }
        )
    version = f"sha256:{digest.hexdigest()[:16]}" if any(file["exists"] for file in files) else "not available"
    audit = latest_production_readiness_gate()
    audit_metrics = (audit or {}).get("metrics") or {}
    return {
        "version": version,
        "files": files,
        "total_rule_rows": total_rows,
        "approved_rule_rows": approved_rows,
        "review_or_nonapproved_rule_rows": review_rows,
        "latest_audit_timestamp": audit_metrics.get("latest_audit_timestamp") or "not available",
        "latest_audit_git_commit": audit_metrics.get("git_commit") or "not available",
        "latest_audit_code_version": audit_metrics.get("code_version") or "not available",
        "production_readiness_gate": audit,
    }


def tool_version_manifest() -> dict[str, str]:
    return {
        "ncbi_datasets_version": command_version(["datasets", "--version"]),
        "taxonkit_version": command_version(["taxonkit", "version"]),
    }


def iso_from_timestamp(value: float | int | None) -> str:
    if not value:
        return "not recorded"
    return datetime.fromtimestamp(float(value), timezone.utc).isoformat()


def load_priority_pathogens(path: Path | None = None) -> list[dict[str, str]]:
    source = path or Path(__file__).with_name("priority_pathogens.csv")
    if not source.exists():
        return []
    with source.open(newline="", encoding="utf-8") as handle:
        return [
            {str(key): normalize(value) for key, value in row.items() if key is not None}
            for row in csv.DictReader(handle)
        ]


def calculate_taxon_quality(row: TaxonStats) -> dict[str, Any]:
    total = max(row.rows, 1)
    country = percent(row.country_usable, total)
    host_source = percent(row.host_or_source_usable, total)
    year = percent(row.year_usable, total)
    quality = percent(row.quality_available, total)
    isolation = percent(row.isolation_usable, total)
    confidence = percent(row.confidence_available, total)
    top_project_share = percent(row.bioprojects.most_common(1)[0][1], total) if row.bioprojects else 100.0
    diversity = max(0.0, 100.0 - top_project_share)
    score = round(
        0.20 * country
        + 0.20 * host_source
        + 0.15 * year
        + 0.15 * quality
        + 0.10 * diversity
        + 0.10 * isolation
        + 0.10 * confidence,
        2,
    )
    if score >= 85:
        grade = "Excellent"
    elif score >= 70:
        grade = "Good"
    elif score >= 50:
        grade = "Moderate"
    elif score >= 25:
        grade = "Poor"
    else:
        grade = "Very poor"
    return {
        "taxon": row.name,
        "rank": row.rank,
        "assemblies": row.rows,
        "metadata_quality_score": score,
        "metadata_quality_grade": grade,
        "country_completeness_percent": country,
        "host_or_source_completeness_percent": host_source,
        "collection_year_completeness_percent": year,
        "assembly_quality_availability_percent": quality,
        "bioproject_diversity_percent": round(diversity, 2),
        "isolation_source_completeness_percent": isolation,
        "standardization_confidence_percent": confidence,
        "denominator": row.rows,
        "denominator_note": denominator_note(row.rows, f"assemblies assigned to {row.name}"),
    }


def build_narrative(summary: dict[str, Any]) -> dict[str, str]:
    overview = summary["overview"]
    snapshot_id = summary.get("snapshot_id", "unknown snapshot")
    generated_at = summary.get("generated_at", "unknown date")
    top_genera = summary["taxonomic_landscape"].get("top_genera", [])
    top_species = summary["taxonomic_landscape"].get("top_species", [])
    top_countries = summary.get("geographic_bias", {}).get("countries", [])
    host_categories = summary.get("host_source_bias", {}).get("host_categories", [])
    assembly_levels = summary.get("assembly_quality", {}).get("assembly_levels", [])
    completeness = {row["field"]: row for row in summary["metadata_completeness"]}
    country = completeness.get("Country", {})
    host = completeness.get("Host", {})
    source = completeness.get("Isolation source", {})
    sample = completeness.get("Sample type", {})
    env = completeness.get("Environment", {})
    year = completeness.get("Collection year from collection-date metadata", {}) or completeness.get("Collection year", {})
    top10_share = summary["taxonomic_landscape"].get("top_10_genus_share_percent", 0)
    denominator = int(overview.get("unique_assemblies") or 0)
    top_genus_names = ", ".join(row["label"] for row in top_genera[:5]) or "the most represented genera"
    top_species_names = ", ".join(row["label"] for row in top_species[:5]) or "the most represented species labels"
    top_country = top_countries[0] if top_countries else {"label": "not available", "percent": 0, "count": 0}
    top_host_category = host_categories[0] if host_categories else {"label": "not available", "percent": 0, "count": 0}
    top_assembly = assembly_levels[0] if assembly_levels else {"label": "not available", "percent": 0, "count": 0}
    canonical_species_pages = int(overview.get("canonical_species_pages") or 0)
    unresolved_species_pages = int(overview.get("unresolved_species_level_label_pages") or 0)
    provisional_species_pages = int(overview.get("provisional_taxonomic_label_pages") or 0)
    noncanonical_species_pages = int(overview.get("noncanonical_species_label_pages") or 0)

    abstract = (
        f"FetchM Global Metadata Insights snapshot {snapshot_id} ({generated_at}) summarized {denominator:,} unique public bacterial "
        f"genome assemblies from {overview['metadata_files_scanned']:,} ready metadata files, representing {overview['species_observed']:,} "
        f"observed species labels, {overview['genera_observed']:,} genera, and {overview['bioprojects_observed']:,} BioProjects. "
        f"The live taxon catalog separated {canonical_species_pages:,} canonical species pages from {unresolved_species_pages:,} "
        f"unresolved species-level labels, {provisional_species_pages:,} provisional labels, and {noncanonical_species_pages:,} non-canonical labels. "
        f"Repository representation was uneven: the ten most represented genera accounted for {top10_share}% of assemblies, and the "
        f"largest standardized country category was {top_country['label']} ({top_country['count']:,} assemblies; {top_country['percent']}%). "
        "These results describe public repository composition rather than true bacterial prevalence."
    )
    dataset_overview = (
        f"Dataset overview. Snapshot {snapshot_id} included {denominator:,} non-redundant assemblies after removing "
        f"{overview['duplicate_rows_skipped']:,} duplicate metadata rows by Assembly Accession. The scan covered "
        f"{overview['metadata_rows_scanned']:,} metadata rows from {overview['metadata_files_scanned']:,} files and skipped "
        f"{overview['metadata_files_skipped']:,} unavailable files."
    )
    completeness_text = (
        f"Metadata completeness and standardization. Standardized country assignments were available for "
        f"{country.get('standardized_usable', 0):,}/{denominator:,} assemblies ({country.get('standardized_usable_percent', 0)}%), "
        f"standardized host assignments for {host.get('standardized_usable', 0):,} ({host.get('standardized_usable_percent', 0)}%), "
        f"standardized isolation-source assignments for {source.get('standardized_usable', 0):,} ({source.get('standardized_usable_percent', 0)}%), "
        f"sample-type assignments for {sample.get('standardized_usable', 0):,} ({sample.get('standardized_usable_percent', 0)}%), "
        f"environment assignments for {env.get('standardized_usable', 0):,} ({env.get('standardized_usable_percent', 0)}%), and usable "
        f"collection years standardized from collection-date metadata for {year.get('standardized_usable', 0):,} ({year.get('standardized_usable_percent', 0)}%). Standardized-only recovered records included "
        f"{country.get('standardized_only_rescued', country.get('rescued_records', 0)):,} country assignments, "
        f"{host.get('standardized_only_rescued', host.get('rescued_records', 0)):,} host assignments, and "
        f"{source.get('standardized_only_rescued', source.get('rescued_records', 0)):,} isolation-source assignments. Raw-field presence and "
        "standardized-field availability are intentionally reported separately because present free text may remain too vague, inconsistent, "
        "or unmapped for reliable filtering."
    )
    taxonomic_text = (
        f"Taxonomic concentration. The most represented genera were {top_genus_names}, and the most represented species labels were "
        f"{top_species_names}. The top ten genera accounted for {top10_share}% of all non-redundant assemblies, indicating that public "
        "bacterial genome repositories are dominated by a limited set of heavily sequenced taxa. FetchM keeps unresolved and provisional "
        "species-level labels searchable for comprehensive access, but reports them separately from canonical binomial species labels."
    )
    geography_text = (
        f"Geographic representation. Standardized country metadata identified {overview['countries_observed']:,} country categories. "
        f"The largest country category was {top_country['label']}, with {top_country['count']:,} assemblies ({top_country['percent']}% of "
        "the full assembly denominator). These values measure representation in public repositories and should not be interpreted as "
        "country-level disease burden or environmental abundance."
    )
    host_text = (
        f"Host and source representation. Host/source categorization was led by {top_host_category['label']} "
        f"({top_host_category['count']:,} assemblies; {top_host_category['percent']}%). Missing or ambiguous values are separated from "
        "biological host categories so that absent submitter fields are not treated as host dominance."
    )
    temporal_quality_text = (
        f"Temporal growth and assembly quality. Release-year and collection-year fields were used to summarize temporal coverage, while "
        f"assembly-level metadata showed {top_assembly['label']} as the largest assembly-level category "
        f"({top_assembly['count']:,} assemblies; {top_assembly['percent']}%). Assembly-quality availability was incorporated into the "
        "metadata quality score when completeness, contamination, N50, contig count, genome size, or GC fields were present."
    )
    quality_text = (
        f"Metadata quality and readiness. Under the default readiness rule, {overview['qc_ready_taxa']:,} taxa had at least 100 assemblies, "
        "at least 70% standardized country completeness, at least 70% host/source completeness, and at least 50% collection-year completeness. "
        "The metadata quality score combines country, host/source, collection year, assembly-quality availability, BioProject diversity, "
        "isolation-source completeness, and standardization-confidence components."
    )
    bias_text = (
        "BioProject dominance and sampling bias. Bias warnings quantify repository sampling structure, including top BioProject share, "
        "country dominance, host dominance, and collection-year dominance. Missing values such as absent, unknown, not collected, or not "
        "applicable are excluded from biological dominance warnings and should instead be interpreted as metadata missingness."
    )
    limitations = (
        "Several limitations should be considered. FetchM Global Insights reflects publicly available genome metadata and is influenced by "
        "database submission practices, surveillance priorities, sequencing capacity, outbreak sampling, and reporting standards. Standardization "
        "can recover many useful records but cannot infer reliable biological context from unsupported or ambiguous submitter text. All percentages "
        "therefore describe public repository representation, not true global bacterial abundance, disease burden, or environmental prevalence."
    )
    methods = (
        f"Global Metadata Insights snapshot {snapshot_id} was generated from ready FetchM standardized metadata files. Unique assemblies were "
        "counted by Assembly Accession, with species-level rows preferred over genus-level rows and newer synced taxa scanned first within each rank. "
        "Metadata completeness was calculated separately for raw submitter fields and standardized analysis fields. Empty, absent, unknown, not "
        "collected, not applicable, unidentified, and similarly non-informative values were treated as unusable metadata. Species-level labels "
        "were classified as canonical binomial species, unresolved placeholder labels, provisional Candidatus labels, or non-canonical labels; all "
        "remain searchable, but only canonical binomials are counted as canonical species."
    )
    figure_legend = (
        f"Figure X. Repository representation and metadata completeness of FetchM-indexed bacterial genome assemblies. Values were calculated "
        f"from Global Metadata Insights snapshot {snapshot_id} using {denominator:,} non-redundant public bacterial assemblies. Panels summarize "
        "taxonomic concentration, geographic representation, host/source categories, temporal coverage, assembly-level metadata, raw-versus-"
        "standardized metadata availability, and automated bias warnings. Source data are provided in the snapshot tables."
    )
    table_caption = (
        f"Table X. Metadata completeness, standardization impact, metadata quality scores, QC-ready taxa, and repository bias warnings for "
        f"FetchM Global Metadata Insights snapshot {snapshot_id}. Denominators use {denominator:,} non-redundant assemblies unless otherwise stated."
    )
    results = " ".join([
        dataset_overview,
        completeness_text,
        taxonomic_text,
        geography_text,
        host_text,
        temporal_quality_text,
        quality_text,
        bias_text,
    ])
    return {
        "abstract": abstract,
        "results": results,
        "dataset_overview": dataset_overview,
        "metadata_completeness": completeness_text,
        "taxonomic_concentration": taxonomic_text,
        "geographic_representation": geography_text,
        "host_source_representation": host_text,
        "temporal_growth_and_assembly_quality": temporal_quality_text,
        "metadata_quality_by_taxon": quality_text,
        "bioproject_dominance_and_sampling_bias": bias_text,
        "methods": methods,
        "limitations": limitations,
        "figure_legend": figure_legend,
        "table_caption": table_caption,
    }


def _figure_label(value: Any, limit: int = 34) -> str:
    text = normalize(value)
    if len(text) <= limit:
        return text
    return text[: limit - 1].rstrip() + "…"


def _write_source_data(path: Path, rows: list[dict[str, Any]]) -> None:
    if not rows:
        write_csv(path, [], ["label", "value"])
        return
    fieldnames: list[str] = []
    for row in rows:
        for key in row.keys():
            if key not in fieldnames:
                fieldnames.append(key)
    write_csv(path, rows, fieldnames)


def _zip_relative_files(base_dir: Path, zip_path: Path, relative_paths: list[str]) -> None:
    zip_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for relative_path in relative_paths:
            source = base_dir / relative_path
            if source.exists() and source.is_file():
                archive.write(source, relative_path)


def _write_minimal_docx(path: Path, title: str, sections: list[tuple[str, str]]) -> None:
    """Write a small valid DOCX without requiring python-docx."""
    from xml.sax.saxutils import escape

    def paragraph(text: str, *, heading: bool = False) -> str:
        style = '<w:pStyle w:val="Heading1"/>' if heading else ""
        return (
            "<w:p><w:pPr>" + style + "</w:pPr><w:r><w:t xml:space=\"preserve\">"
            + escape(text)
            + "</w:t></w:r></w:p>"
        )

    body = [paragraph(title, heading=True)]
    for section_title, section_body in sections:
        body.append(paragraph(section_title, heading=True))
        for chunk in textwrap.wrap(section_body or "Not available.", width=100) or [""]:
            body.append(paragraph(chunk))
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:body>'
        + "".join(body)
        + '<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>'
        + '</w:body></w:document>'
    )
    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '</Types>'
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        '</Relationships>'
    )
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", rels)
        archive.writestr("word/document.xml", document_xml)


def _save_figure(fig: Any, base_dir: Path, stem: str) -> dict[str, str]:
    figure_dir = base_dir / "figures"
    figure_dir.mkdir(parents=True, exist_ok=True)
    files = {
        "svg": f"figures/{stem}.svg",
        "pdf": f"figures/{stem}.pdf",
        "png": f"figures/{stem}.png",
    }
    for ext, relative_path in files.items():
        kwargs = {"bbox_inches": "tight", "facecolor": "white"}
        if ext == "png":
            kwargs["dpi"] = 600
        fig.savefig(base_dir / relative_path, **kwargs)
    return files


def _plot_barh(ax: Any, rows: list[dict[str, Any]], *, label_key: str = "label", value_key: str = "count", title: str = "", color: str = "#165c4e") -> None:
    rows = [row for row in rows if row.get(value_key) not in (None, "")]
    if not rows:
        ax.text(0.5, 0.5, "No data available", ha="center", va="center", transform=ax.transAxes)
        ax.set_axis_off()
        return
    labels = [_figure_label(row.get(label_key, ""), 32) for row in rows][::-1]
    values = [float(row.get(value_key) or 0) for row in rows][::-1]
    ax.barh(labels, values, color=color, alpha=0.92)
    ax.set_title(title, loc="left", fontweight="bold")
    ax.grid(axis="x", color="#d9ded6", linewidth=0.7, alpha=0.65)
    ax.spines[["top", "right", "left"]].set_visible(False)
    ax.tick_params(axis="y", length=0)


def generate_publication_exports(summary: dict[str, Any], snapshot_dir: Path, table_dir: Path) -> None:
    """Generate manuscript-grade static figures, source data, and reports for a Global Insights snapshot."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_pdf import PdfPages
    except Exception as exc:  # pragma: no cover - depends on deployment image
        summary.setdefault("qa", {}).setdefault("checks", []).append(
            {"check": "publication_exports_generated", "status": "warning", "detail": f"Matplotlib export unavailable: {exc}"}
        )
        summary.setdefault("qa", {}).setdefault("manuscript_readiness", {})["all_figures_exported"] = False
        return

    figure_dir = snapshot_dir / "figures"
    source_dir = snapshot_dir / "source_data"
    report_dir = snapshot_dir / "reports"
    for directory in (figure_dir, source_dir, report_dir):
        directory.mkdir(parents=True, exist_ok=True)

    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "axes.edgecolor": "#52615b",
            "axes.labelcolor": "#26332f",
            "xtick.color": "#52615b",
            "ytick.color": "#26332f",
            "figure.facecolor": "white",
            "axes.facecolor": "white",
        }
    )

    overview = summary.get("overview", {})
    manuscript = summary.get("manuscript", {})
    snapshot_id = summary.get("snapshot_id", "unknown snapshot")
    generated_at = summary.get("generated_at", "unknown date")
    denominator = int(overview.get("unique_assemblies") or 0)
    figure_exports: list[dict[str, Any]] = []
    all_figure_files: list[str] = []
    source_files: list[str] = []

    def add_figure(fig: Any, stem: str, title: str, rows: list[dict[str, Any]], source_name: str, interpretation: str) -> None:
        files = _save_figure(fig, snapshot_dir, stem)
        plt.close(fig)
        source_relative = f"source_data/{source_name}.csv"
        source_table_name = source_relative
        figure_number = len(figure_exports) + 1
        interpretation_note = interpretation or "This figure summarizes repository representation and should not be interpreted as true biological prevalence."
        legend = (
            f"Figure {figure_number}. {title}. Values were calculated from Global Metadata Insights snapshot {snapshot_id} "
            f"generated at {generated_at}. The exact denominator is {denominator:,} non-redundant public bacterial assemblies unless the panel label states otherwise. "
            f"Source table: {source_table_name}. {interpretation_note}"
        )
        legend_relative = f"figures/{stem}_legend.txt"
        (snapshot_dir / legend_relative).write_text(legend + "\n", encoding="utf-8")
        source_rows = []
        for row in rows:
            enriched = dict(row)
            enriched.setdefault("snapshot_id", snapshot_id)
            enriched.setdefault("figure_id", stem)
            enriched.setdefault("denominator", denominator)
            enriched.setdefault("denominator_note", denominator_note(denominator))
            source_rows.append(enriched)
        _write_source_data(snapshot_dir / source_relative, source_rows)
        all_figure_files.extend([*files.values(), legend_relative])
        source_files.append(source_relative)
        figure_exports.append(
            {
                "figure_id": stem,
                "title": title,
                "files": files,
                "legend_file": legend_relative,
                "source_data": source_relative,
                "source_table_name": source_table_name,
                "denominator": denominator,
                "denominator_note": denominator_note(denominator),
                "snapshot_id": snapshot_id,
                "generated_at": generated_at,
                "interpretation_note": interpretation_note,
                "legend": legend,
            }
        )

    # Figure 1: global snapshot overview.
    overview_rows = [
        {"metric": "Unique assemblies", "value": overview.get("unique_assemblies", 0)},
        {"metric": "Observed species labels", "value": overview.get("species_observed", 0)},
        {"metric": "Observed genera", "value": overview.get("genera_observed", 0)},
        {"metric": "BioProjects", "value": overview.get("bioprojects_observed", 0)},
        {"metric": "Ready metadata files", "value": overview.get("metadata_files_scanned", 0)},
        {"metric": "QC-ready taxa", "value": overview.get("qc_ready_taxa", 0)},
    ]
    fig, ax = plt.subplots(figsize=(11, 6.2))
    ax.set_axis_off()
    ax.text(0.02, 0.93, "Global Metadata Insights", fontsize=21, fontweight="bold", color="#165c4e", transform=ax.transAxes)
    ax.text(0.02, 0.86, f"Snapshot {snapshot_id} · {generated_at}", fontsize=10.5, color="#52615b", transform=ax.transAxes)
    positions = [(0.03, 0.58), (0.36, 0.58), (0.69, 0.58), (0.03, 0.28), (0.36, 0.28), (0.69, 0.28)]
    for row, (x, y) in zip(overview_rows, positions):
        ax.add_patch(plt.Rectangle((x, y), 0.28, 0.2, transform=ax.transAxes, facecolor="#f5efe1", edgecolor="#cbd8d2", linewidth=1.2))
        ax.text(x + 0.025, y + 0.115, f"{int(row['value']):,}", fontsize=20, fontweight="bold", color="#165c4e", transform=ax.transAxes)
        ax.text(x + 0.025, y + 0.055, row["metric"], fontsize=10.5, color="#52615b", transform=ax.transAxes)
    add_figure(fig, "figure_1_global_snapshot", "Global snapshot overview", overview_rows, "figure_1_global_snapshot_source_data", manuscript.get("figure_legend", ""))

    # Figure 2: raw vs standardized metadata completeness.
    completeness_rows = summary.get("metadata_completeness", [])
    fig, ax = plt.subplots(figsize=(11, 6.4))
    fields = [row.get("field", "") for row in completeness_rows]
    x = range(len(fields))
    raw_values = [float(row.get("raw_usable_percent") or 0) for row in completeness_rows]
    std_values = [float(row.get("standardized_usable_percent") or 0) for row in completeness_rows]
    width = 0.36
    ax.bar([i - width / 2 for i in x], raw_values, width=width, label="Raw field present", color="#b8844d")
    ax.bar([i + width / 2 for i in x], std_values, width=width, label="Standardized", color="#165c4e")
    ax.set_xticks(list(x), [_figure_label(field, 20) for field in fields], rotation=25, ha="right")
    ax.set_ylabel("Assemblies (%)")
    ax.set_ylim(0, 100)
    ax.set_title("Raw versus standardized metadata completeness", loc="left", fontweight="bold")
    ax.legend(frameon=False)
    ax.grid(axis="y", color="#d9ded6", linewidth=0.7, alpha=0.65)
    ax.spines[["top", "right"]].set_visible(False)
    add_figure(fig, "figure_2_metadata_completeness", "Raw versus standardized metadata completeness", completeness_rows, "figure_2_metadata_completeness_source_data", manuscript.get("metadata_completeness", ""))

    # Figure 3: top genera and species.
    fig, axes = plt.subplots(1, 2, figsize=(13, 7))
    top_genera = summary.get("taxonomic_landscape", {}).get("top_genera", [])[:12]
    top_species = summary.get("taxonomic_landscape", {}).get("top_species", [])[:12]
    _plot_barh(axes[0], top_genera, title="Top genera", color="#165c4e")
    _plot_barh(axes[1], top_species, title="Top species labels", color="#8a5a16")
    fig.suptitle("Taxonomic concentration in public bacterial genome repositories", x=0.03, ha="left", fontweight="bold", fontsize=15)
    add_figure(fig, "figure_3_taxonomic_concentration", "Taxonomic concentration", top_genera + top_species, "figure_3_taxonomic_concentration_source_data", manuscript.get("taxonomic_concentration", ""))

    # Figure 4: geography.
    country_rows = summary.get("geographic_bias", {}).get("countries", [])[:20]
    fig, ax = plt.subplots(figsize=(10.5, 7.2))
    _plot_barh(ax, country_rows, title="Top standardized countries", color="#165c4e")
    ax.set_xlabel("Assemblies")
    add_figure(fig, "figure_4_geographic_representation", "Geographic representation", country_rows, "figure_4_geographic_representation_source_data", manuscript.get("geographic_representation", ""))

    # Figure 5: host/source representation.
    fig, axes = plt.subplots(1, 2, figsize=(13, 7))
    host_categories = summary.get("host_source_bias", {}).get("host_categories", [])[:12]
    top_sources = summary.get("host_source_bias", {}).get("sources", [])[:12]
    _plot_barh(axes[0], host_categories, title="Host/source categories", color="#165c4e")
    _plot_barh(axes[1], top_sources, title="Top standardized isolation sources", color="#8a5a16")
    fig.suptitle("Host, source, and environment representation", x=0.03, ha="left", fontweight="bold", fontsize=15)
    add_figure(fig, "figure_5_host_source_representation", "Host and source representation", host_categories + top_sources, "figure_5_host_source_representation_source_data", manuscript.get("host_source_representation", ""))

    # Figure 6: temporal growth and assembly levels.
    fig, axes = plt.subplots(1, 2, figsize=(13, 6.6))
    yearly = summary.get("yearly_growth", [])[-20:]
    if yearly:
        axes[0].plot([row.get("year") for row in yearly], [int(row.get("assemblies") or 0) for row in yearly], color="#165c4e", linewidth=2.4)
        axes[0].fill_between([row.get("year") for row in yearly], [int(row.get("assemblies") or 0) for row in yearly], color="#165c4e", alpha=0.15)
        axes[0].tick_params(axis="x", rotation=45)
        axes[0].set_title("Annual genome availability", loc="left", fontweight="bold")
        axes[0].set_ylabel("Assemblies")
        axes[0].grid(axis="y", color="#d9ded6", linewidth=0.7, alpha=0.65)
        axes[0].spines[["top", "right"]].set_visible(False)
    else:
        axes[0].text(0.5, 0.5, "No year data", ha="center", va="center")
        axes[0].set_axis_off()
    assembly_rows = summary.get("assembly_quality", {}).get("assembly_levels", [])[:10]
    _plot_barh(axes[1], assembly_rows, title="Assembly levels", color="#8a5a16")
    add_figure(fig, "figure_6_temporal_growth_assembly_levels", "Temporal growth and assembly levels", yearly + assembly_rows, "figure_6_temporal_growth_source_data", manuscript.get("temporal_growth_and_assembly_quality", ""))

    # Figure 7: metadata quality and QC-ready taxa.
    quality_rows = summary.get("metadata_quality", [])[:20]
    fig, ax = plt.subplots(figsize=(11, 8))
    plot_rows = [{"label": row.get("taxon"), "count": row.get("metadata_quality_score", 0)} for row in quality_rows]
    _plot_barh(ax, plot_rows, title="Top metadata quality scores", color="#165c4e")
    ax.set_xlabel("Metadata quality score (0-100)")
    ax.set_xlim(0, 100)
    add_figure(fig, "figure_7_metadata_quality_qc_ready", "Metadata quality and QC-ready taxa", quality_rows, "figure_7_metadata_quality_source_data", manuscript.get("metadata_quality_by_taxon", ""))

    # Figure 8: bias warning counts.
    warning_rows = summary.get("bias_warnings", [])
    warning_counter = Counter(row.get("bias_type", "Unknown") for row in warning_rows)
    warning_plot_rows = [{"label": label, "count": count} for label, count in warning_counter.most_common(15)]
    fig, ax = plt.subplots(figsize=(10.5, 7))
    _plot_barh(ax, warning_plot_rows, title="Automated bias warnings by type", color="#8a5a16")
    ax.set_xlabel("Warnings")
    add_figure(fig, "figure_8_bias_warnings", "Automated bias warnings", warning_plot_rows, "figure_8_bias_warnings_source_data", manuscript.get("bioproject_dominance_and_sampling_bias", ""))

    # Figure 9: standardization impact/corrections.
    correction_rows = summary.get("standardization_impact", {}).get("top_corrections", [])[:18]
    correction_plot_rows = [
        {
            "label": f"{row.get('field')}: {_figure_label(row.get('standardized_value'), 24)}",
            "count": int(row.get("records_rescued") or 0),
            **row,
        }
        for row in correction_rows
    ]
    fig, ax = plt.subplots(figsize=(11, 8))
    _plot_barh(ax, correction_plot_rows, title="Most common raw-to-standardized mappings", color="#165c4e")
    ax.set_xlabel("Mapped records")
    add_figure(fig, "figure_9_standardization_impact", "Standardization impact", correction_plot_rows, "figure_9_standardization_impact_source_data", manuscript.get("metadata_completeness", ""))

    # Manuscript report files.
    sections = [
        ("Abstract-style summary", manuscript.get("abstract", "")),
        ("Dataset overview", manuscript.get("dataset_overview", "")),
        ("Metadata completeness and standardization", manuscript.get("metadata_completeness", "")),
        ("Taxonomic concentration", manuscript.get("taxonomic_concentration", "")),
        ("Geographic representation", manuscript.get("geographic_representation", "")),
        ("Host/source representation", manuscript.get("host_source_representation", "")),
        ("Temporal growth and assembly quality", manuscript.get("temporal_growth_and_assembly_quality", "")),
        ("Metadata quality by taxon", manuscript.get("metadata_quality_by_taxon", "")),
        ("BioProject dominance and sampling bias", manuscript.get("bioproject_dominance_and_sampling_bias", "")),
        ("Methods", manuscript.get("methods", "")),
        ("Limitations", manuscript.get("limitations", "")),
        ("Figure legend", manuscript.get("figure_legend", "")),
        ("Table caption", manuscript.get("table_caption", "")),
    ]
    report_md = report_dir / "global_insights_report.md"
    md_lines = [f"# Global Metadata Insights Report", "", f"Snapshot: `{snapshot_id}`", f"Generated: `{generated_at}`", ""]
    for title, body in sections:
        md_lines.extend([f"## {title}", "", body or "Not available.", ""])
    md_lines.extend(["## Publication Figures", ""])
    for figure in figure_exports:
        md_lines.extend([f"### {figure['title']}", "", figure["legend"] or "Legend not available.", ""])
    report_md.write_text("\n".join(md_lines), encoding="utf-8")

    report_docx = report_dir / "global_insights_report.docx"
    docx_sections = sections + [(figure["title"], figure.get("legend", "")) for figure in figure_exports]
    try:
        from docx import Document
        document = Document()
        document.add_heading("Global Metadata Insights Report", level=1)
        document.add_paragraph(f"Snapshot: {snapshot_id}")
        document.add_paragraph(f"Generated: {generated_at}")
        for title, body in docx_sections:
            document.add_heading(title, level=2)
            document.add_paragraph(body or "Not available.")
        document.save(report_docx)
    except Exception as exc:  # pragma: no cover - optional export dependency
        _write_minimal_docx(
            report_docx,
            "Global Metadata Insights Report",
            [("Snapshot", f"Snapshot: {snapshot_id}\nGenerated: {generated_at}"), *docx_sections],
        )
        summary.setdefault("qa", {}).setdefault("checks", []).append(
            {"check": "docx_report_generated", "status": "pass", "detail": f"DOCX generated with built-in fallback because python-docx was unavailable: {exc}"}
        )

    report_pdf = report_dir / "global_insights_report.pdf"
    try:
        with PdfPages(report_pdf) as pdf:
            for title, body in sections:
                page = plt.figure(figsize=(8.27, 11.69))
                page.text(0.08, 0.94, title, fontsize=15, fontweight="bold", color="#165c4e")
                wrapped = "\n".join(textwrap.wrap(body or "Not available.", width=94))
                page.text(0.08, 0.89, wrapped, fontsize=9.5, color="#26332f", va="top")
                page.text(0.08, 0.04, f"Snapshot {snapshot_id}", fontsize=8, color="#52615b")
                pdf.savefig(page, bbox_inches="tight")
                plt.close(page)
    except Exception as exc:  # pragma: no cover - optional export backend
        report_pdf = None
        summary.setdefault("qa", {}).setdefault("checks", []).append(
            {"check": "pdf_report_generated", "status": "warning", "detail": f"PDF report export unavailable: {exc}"}
        )

    case_files: list[str] = []
    case_dir = report_dir / "case_studies"
    case_dir.mkdir(parents=True, exist_ok=True)
    for case in summary.get("case_studies") or []:
        slug = safe_slug(str(case.get("taxon") or "case-study"))
        relative = f"reports/case_studies/{slug}.md"
        lines = [
            f"# {case.get('taxon', 'Case study')}",
            "",
            f"Status: {case.get('status', 'available')}",
            f"Assemblies: {case.get('assemblies', 0)}",
            f"Metadata quality: {case.get('metadata_quality_score', 'not available')} ({case.get('metadata_quality_grade', 'not available')})",
            f"Readiness tier: {case.get('metadata_readiness_tier', 'not available')}",
            f"Top countries: {case.get('top_countries', 'not available')}",
            f"Top hosts: {case.get('top_hosts', 'not available')}",
            f"Top sources: {case.get('top_sources', 'not available')}",
            f"Top BioProject: {case.get('top_bioproject', 'not recorded')} ({case.get('top_bioproject_share_percent', 0)}%)",
            "",
            f"Sampling caution: {case.get('sampling_caution', 'not available')}",
            "",
            str(case.get("denominator_note", "")),
        ]
        (snapshot_dir / relative).write_text("\n".join(lines) + "\n", encoding="utf-8")
        case_files.append(relative)

    table_files = [f"tables/{path.name}" for path in table_dir.glob("*.csv")]
    figure_zip = snapshot_dir / "global_insights_figures.zip"
    table_zip = snapshot_dir / "global_insights_tables.zip"
    source_zip = snapshot_dir / "source_data_for_figures.zip"
    case_zip = snapshot_dir / "global_insights_case_studies.zip"
    _zip_relative_files(snapshot_dir, figure_zip, all_figure_files)
    _zip_relative_files(snapshot_dir, table_zip, table_files)
    _zip_relative_files(snapshot_dir, source_zip, source_files)
    _zip_relative_files(snapshot_dir, case_zip, case_files)

    report_downloads = {
        "global_insights_report_md": "reports/global_insights_report.md",
        "global_insights_figures_zip": "global_insights_figures.zip",
        "global_insights_tables_zip": "global_insights_tables.zip",
        "source_data_for_figures_zip": "source_data_for_figures.zip",
        "global_insights_case_studies_zip": "global_insights_case_studies.zip",
    }
    if report_docx:
        report_downloads["global_insights_report_docx"] = "reports/global_insights_report.docx"
    if report_pdf:
        report_downloads["global_insights_report_pdf"] = "reports/global_insights_report.pdf"
    summary.setdefault("downloads", {}).update(report_downloads)
    summary["figure_exports"] = figure_exports
    audit = write_publication_export_audit(summary, snapshot_dir)
    summary.setdefault("qa", {}).setdefault("checks", []).append(
        {"check": "publication_exports_generated", "status": "pass", "detail": f"{len(figure_exports)} figures exported as SVG, PDF, and 600-dpi PNG with source data."}
    )
    summary.setdefault("qa", {}).setdefault("checks", []).append(
        {"check": "publication_export_audit", "status": audit.get("status", "warning"), "detail": f"{audit.get('figure_count', 0)} figures audited for SVG/PDF/PNG/source CSV/legend TXT/denominator artifacts."}
    )
    summary.setdefault("qa", {}).setdefault("manuscript_readiness", {})["all_figures_exported"] = len(figure_exports) >= 9
    summary.setdefault("qa", {}).setdefault("manuscript_readiness", {})["publication_export_audit_passed"] = audit.get("status") == "pass"

def generate_global_insights_snapshot(
    taxa: Iterable[dict[str, Any] | TaxonInput],
    output_root: Path,
    *,
    app_version: str,
    app_commit: str,
    snapshot_id: str | None = None,
    demo: bool = False,
    canonical_root_source: bool = False,
    source_snapshot_id: str | None = None,
) -> dict[str, Any]:
    if demo:
        return generate_demo_snapshot(output_root, app_version=app_version, app_commit=app_commit, snapshot_id=snapshot_id)

    snapshot_id = snapshot_id or f"{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')}_global_insights"
    output_root.mkdir(parents=True, exist_ok=True)
    snapshot_dir = output_root / "snapshots" / safe_slug(snapshot_id)
    tmp_dir = output_root / "snapshots" / f".{safe_slug(snapshot_id)}.tmp"
    if tmp_dir.exists():
        shutil.rmtree(tmp_dir)
    tmp_dir.mkdir(parents=True, exist_ok=True)
    table_dir = tmp_dir / "tables"
    table_dir.mkdir(parents=True, exist_ok=True)

    status_path = tmp_dir / "status.json"
    status_path.write_text(json.dumps({"snapshot_id": snapshot_id, "status": "running", "started_at": utc_now()}, indent=2), encoding="utf-8")

    seen_accessions: set[str] = set()
    unique_total = 0
    duplicate_rows = 0
    metadata_rows_scanned = 0
    files_scanned = 0
    files_skipped = 0
    latest_metadata_mtime = 0.0
    latest_taxon_synced_at = ""
    genus_counter: Counter[str] = Counter()
    species_counter: Counter[str] = Counter()
    species_label_class_assembly_counter: Counter[str] = Counter()
    species_label_class_names: dict[str, set[str]] = defaultdict(set)
    taxon_label_class_counter: Counter[str] = Counter()
    taxon_label_class_names: dict[str, set[str]] = defaultdict(set)
    country_counter: Counter[str] = Counter()
    continent_counter: Counter[str] = Counter()
    subcontinent_counter: Counter[str] = Counter()
    non_country_geography_counter: Counter[str] = Counter()
    host_counter: Counter[str] = Counter()
    source_counter: Counter[str] = Counter()
    host_category_counter: Counter[str] = Counter()
    assembly_level_counter: Counter[str] = Counter()
    year_counter: Counter[str] = Counter()
    cumulative_year_counter: Counter[str] = Counter()
    bioproject_counter: Counter[str] = Counter()
    bioproject_taxon_counter: dict[str, Counter[str]] = defaultdict(Counter)
    confidence_counter: Counter[str] = Counter()
    field_confidence_counters: dict[str, Counter[str]] = defaultdict(Counter)
    correction_counter: Counter[tuple[str, str, str, str, str, str]] = Counter()
    taxon_stats: dict[str, TaxonStats] = {}
    completeness_fields = {
        "Country": {"raw_usable": 0, "standardized_usable": 0, "both_usable": 0, "standardized_only": 0, "raw_only": 0, "changed_mappings": 0},
        "Host": {"raw_usable": 0, "standardized_usable": 0, "both_usable": 0, "standardized_only": 0, "raw_only": 0, "changed_mappings": 0},
        "Collection year from collection-date metadata": {"raw_usable": 0, "standardized_usable": 0, "both_usable": 0, "standardized_only": 0, "raw_only": 0, "changed_mappings": 0},
        "Isolation source": {"raw_usable": 0, "standardized_usable": 0, "both_usable": 0, "standardized_only": 0, "raw_only": 0, "changed_mappings": 0},
        "Sample type": {"raw_usable": 0, "standardized_usable": 0, "both_usable": 0, "standardized_only": 0, "raw_only": 0, "changed_mappings": 0},
        "Environment": {"raw_usable": 0, "standardized_usable": 0, "both_usable": 0, "standardized_only": 0, "raw_only": 0, "changed_mappings": 0},
    }

    simulator_path = table_dir / "simulator_records.csv"
    simulator_fields = [
        "assembly_accession",
        "genus",
        "species",
        "taxon",
        "taxon_rank",
        "raw_country",
        "standardized_country",
        "raw_host",
        "standardized_host",
        "raw_source",
        "standardized_source",
        "raw_collection_date",
        "collection_year",
        "collection_date_source",
        "collection_date_evidence",
        "collection_date_recovery_status",
        "assembly_level",
    ]
    taxon_inputs = sorted_taxa(taxa)
    for taxon in taxon_inputs:
        if taxon.rank == "species":
            label_meta = taxonomy_label_metadata(taxon.name, taxon.rank)
            taxon_label_class_counter[label_meta["key"]] += 1
            taxon_label_class_names[label_meta["key"]].add(taxon.name)

    with simulator_path.open("w", newline="", encoding="utf-8") as simulator_handle:
        simulator_writer = csv.DictWriter(simulator_handle, fieldnames=simulator_fields)
        simulator_writer.writeheader()
        for taxon in taxon_inputs:
            path = resolve_metadata_path(taxon.metadata_clean_path)
            if not path.exists():
                files_skipped += 1
                continue
            files_scanned += 1
            try:
                latest_metadata_mtime = max(latest_metadata_mtime, path.stat().st_mtime)
            except OSError:
                pass
            if taxon.last_synced_at and taxon.last_synced_at > latest_taxon_synced_at:
                latest_taxon_synced_at = taxon.last_synced_at
            taxon_key = f"{taxon.rank}:{taxon.name}"
            default_stat = None if canonical_root_source else taxon_stats.setdefault(taxon_key, TaxonStats(name=taxon.name, rank=taxon.rank))
            with path.open(newline="", encoding="utf-8", errors="replace") as handle:
                reader = csv.DictReader(handle)
                for row in reader:
                    metadata_rows_scanned += 1
                    accession = row_value(row, ASSEMBLY_FIELDS)
                    if not accession:
                        accession = f"{taxon_key}:{metadata_rows_scanned}"
                    if accession in seen_accessions:
                        duplicate_rows += 1
                        continue
                    seen_accessions.add(accession)
                    unique_total += 1

                    organism = row_value(row, ORGANISM_FIELDS)
                    inferred_genus, inferred_species = parse_taxonomy(organism, taxon.name)
                    genus = row_value(row, ("Taxonomy Genus",)) or inferred_genus
                    species = row_value(row, ("Taxonomy Species",)) or inferred_species
                    if canonical_root_source:
                        row_stats = [
                            taxon_stats.setdefault(f"genus:{genus}", TaxonStats(name=genus, rank="genus")),
                            taxon_stats.setdefault(f"species:{species}", TaxonStats(name=species, rank="species")),
                        ]
                    else:
                        row_stats = [default_stat]
                    for stat in row_stats:
                        stat.rows += 1
                    genus_counter[genus] += 1
                    species_counter[species] += 1
                    species_label_meta = taxonomy_label_metadata(species, "species")
                    species_label_class_assembly_counter[species_label_meta["key"]] += 1
                    species_label_class_names[species_label_meta["key"]].add(species)

                    raw_country = row_value(row, COUNTRY_RAW_FIELDS)
                    std_country = row_value(row, COUNTRY_STD_FIELDS)
                    raw_host = row_value(row, HOST_RAW_FIELDS)
                    std_host = row_value(row, HOST_STD_FIELDS)
                    raw_source = row_value(row, SOURCE_RAW_FIELDS)
                    std_source = row_value(row, SOURCE_STD_FIELDS)
                    raw_sample = row_value(row, SAMPLE_RAW_FIELDS)
                    std_sample = row_value(row, SAMPLE_STD_FIELDS)
                    raw_env = row_value(row, ENV_RAW_FIELDS)
                    std_env = row_value(row, ENV_STD_FIELDS)
                    raw_collection_date = row_value(row, ("Collection_Date_Evidence", "BioSample Collection Date", "BioSample Specimen Collection Date", "BioSample Collection Date Remark", "Collection Date"))
                    release_year = parse_year(row_value(row, RELEASE_DATE_FIELDS))
                    collection_year = parse_year(row_value(row, COLLECTION_DATE_FIELDS))
                    assembly_level = row_value(row, ASSEMBLY_LEVEL_FIELDS)
                    bioproject = row_value(row, BIOPROJECT_FIELDS)
                    continent = row_value(row, ("Continent",))
                    subcontinent = row_value(row, ("Subcontinent",))

                    update_field_pair_stats(completeness_fields["Country"], raw_country, std_country)
                    if is_usable(std_country):
                        country_counter[std_country] += 1
                        for stat in row_stats:
                            stat.country_usable += 1
                            stat.countries[std_country] += 1
                    if is_usable(continent):
                        if continent in POLITICAL_CONTINENTS:
                            continent_counter[continent] += 1
                        else:
                            non_country_geography_counter[continent] += 1
                    if is_usable(subcontinent):
                        if subcontinent in POLITICAL_SUBCONTINENTS:
                            subcontinent_counter[subcontinent] += 1
                        else:
                            non_country_geography_counter[subcontinent] += 1
                    update_field_pair_stats(completeness_fields["Host"], raw_host, std_host)
                    if is_usable(std_host):
                        host_counter[std_host] += 1
                        for stat in row_stats:
                            stat.hosts[std_host] += 1
                    update_field_pair_stats(completeness_fields["Isolation source"], raw_source, std_source)
                    if is_usable(std_source):
                        source_counter[std_source] += 1
                        for stat in row_stats:
                            stat.isolation_usable += 1
                            stat.sources[std_source] += 1
                    update_field_pair_stats(completeness_fields["Sample type"], raw_sample, std_sample)
                    update_field_pair_stats(completeness_fields["Environment"], raw_env, std_env)
                    if collection_year:
                        update_field_pair_stats(completeness_fields["Collection year from collection-date metadata"], collection_year, collection_year)
                        for stat in row_stats:
                            stat.year_usable += 1
                            stat.years[collection_year] += 1
                    if release_year:
                        year_counter[release_year] += 1
                    if collection_year:
                        cumulative_year_counter[collection_year] += 1
                    if is_usable(assembly_level):
                        assembly_level_counter[assembly_level] += 1
                    if is_usable(bioproject):
                        bioproject_counter[bioproject] += 1
                        for stat in row_stats:
                            stat.bioprojects[bioproject] += 1
                        bioproject_taxon_counter[bioproject][species if canonical_root_source else taxon.name] += 1
                    if quality_available(row):
                        for stat in row_stats:
                            stat.quality_available += 1
                    if is_usable(std_host) or is_usable(std_source):
                        for stat in row_stats:
                            stat.host_or_source_usable += 1
                    if confidence_available(row):
                        for stat in row_stats:
                            stat.confidence_available += 1
                    for field_name in (*COUNTRY_CONFIDENCE_FIELDS, *HOST_CONFIDENCE_FIELDS):
                        if is_usable(row.get(field_name)):
                            confidence_counter[f"{field_name}: {normalize(row[field_name])}"] += 1

                    host_category_counter[classify_host_source(std_host or raw_host, std_source or raw_source)] += 1
                    for label, raw_value, std_value in (
                        ("Country", raw_country, std_country),
                        ("Host", raw_host, std_host),
                        ("Isolation source", raw_source, std_source),
                        ("Sample type", raw_sample, std_sample),
                        ("Environment", raw_env, std_env),
                        ("Collection year from collection-date metadata", collection_year, collection_year),
                    ):
                        raw_display, evidence, confidence = correction_evidence(row, label, raw_value, std_value)
                        field_confidence_counters[label][confidence_bucket(raw_value, std_value, confidence, evidence)] += 1
                        if not label.startswith("Collection year") and is_usable(std_value) and (not is_usable(raw_value) or normalized_key(raw_value) != normalized_key(std_value)):
                            ctype = correction_type(label, raw_value, std_value, evidence)
                            correction_counter[(label, raw_display, std_value[:120], evidence, confidence, ctype)] += 1

                    simulator_writer.writerow(
                        {
                            "assembly_accession": accession,
                            "genus": genus,
                            "species": species,
                            "taxon": taxon.name,
                            "taxon_rank": taxon.rank,
                            "raw_country": raw_country,
                            "standardized_country": std_country,
                            "raw_host": raw_host,
                            "standardized_host": std_host,
                            "raw_source": raw_source,
                            "standardized_source": std_source,
                            "raw_collection_date": raw_collection_date,
                            "collection_year": collection_year,
                            "collection_date_source": row_value(row, ("Collection_Date_Source",)),
                            "collection_date_evidence": row_value(row, ("Collection_Date_Evidence",)),
                            "collection_date_recovery_status": row_value(row, ("Collection_Date_Recovery_Status",)),
                            "assembly_level": assembly_level,
                        }
                    )

    metadata_completeness = []
    for field_name, counts in completeness_fields.items():
        raw_count = int(counts["raw_usable"])
        std_count = int(counts["standardized_usable"])
        delta_points = round(percent(std_count, unique_total) - percent(raw_count, unique_total), 2)
        metadata_completeness.append(
            {
                "field": field_name,
                "raw_usable": raw_count,
                "raw_usable_percent": percent(raw_count, unique_total),
                "raw_present": raw_count,
                "raw_present_percent": percent(raw_count, unique_total),
                "standardized_usable": std_count,
                "standardized_usable_percent": percent(std_count, unique_total),
                "standardized_records": std_count,
                "standardized_percent": percent(std_count, unique_total),
                "both_usable_records": int(counts["both_usable"]),
                "standardized_only_records": int(counts["standardized_only"]),
                "raw_only_records": int(counts["raw_only"]),
                "raw_not_standardized_records": int(counts["raw_only"]),
                "both_raw_and_standardized_usable": int(counts["both_usable"]),
                "standardized_only_rescued": int(counts["standardized_only"]),
                "raw_only_unresolved": int(counts["raw_only"]),
                "neither_usable": max(0, unique_total - int(counts["both_usable"]) - int(counts["standardized_only"]) - int(counts["raw_only"])),
                "changed_mappings": int(counts["changed_mappings"]),
                "changed_mapping_count": int(counts["changed_mappings"]),
                "raw_values_remapped": int(counts["changed_mappings"]),
                "rescued_records": int(counts["standardized_only"]),
                "gain_percentage_points": delta_points,
                "standardized_gain_percentage_points": delta_points,
                "denominator": unique_total,
                "denominator_note": denominator_note(unique_total),
            }
        )

    taxon_quality = [calculate_taxon_quality(stat) for stat in taxon_stats.values()]
    for quality_row in taxon_quality:
        tier, tier_definition = readiness_tier(quality_row)
        quality_row["metadata_readiness_tier"] = tier
        quality_row["metadata_readiness_definition"] = tier_definition
    taxon_quality.sort(key=lambda row: (-int(row["assemblies"]), -float(row["metadata_quality_score"]), row["taxon"]))
    qc_ready_taxa = [row for row in taxon_quality if row.get("metadata_readiness_tier") in {"strict", "standard"}]
    readiness_counter = Counter(row.get("metadata_readiness_tier", "not-ready") for row in taxon_quality)
    readiness_tiers = [
        {
            "tier": tier,
            "taxa": int(readiness_counter.get(tier, 0)),
            "definition": definition,
            "denominator": len(taxon_quality),
            "denominator_note": denominator_note(len(taxon_quality), "taxa with metadata quality scores"),
        }
        for tier, definition in (
            ("strict", ">=1,000 assemblies, >=80% country, >=80% host/source, >=70% collection year"),
            ("standard", ">=100 assemblies, >=70% country, >=70% host/source, >=50% collection year"),
            ("exploratory", ">=50 assemblies, >=50% country and host/source"),
            ("not-ready", "Below exploratory metadata-readiness thresholds"),
        )
    ]

    bias_warnings: list[dict[str, Any]] = []
    bioproject_dominance: list[dict[str, Any]] = []
    for stat in taxon_stats.values():
        if stat.rows <= 0:
            continue
        top_projects = stat.bioprojects.most_common(10)
        top1 = top_projects[0][1] if top_projects else 0
        top5 = sum(count for _, count in top_projects[:5])
        top10 = sum(count for _, count in top_projects[:10])
        dominance = {
            "taxon": stat.name,
            "rank": stat.rank,
            "assemblies": stat.rows,
            "top_1_bioproject_share_percent": percent(top1, stat.rows),
            "top_5_bioproject_share_percent": percent(top5, stat.rows),
            "top_10_bioproject_share_percent": percent(top10, stat.rows),
            "top_bioproject": top_projects[0][0] if top_projects else "",
            "denominator": stat.rows,
            "denominator_note": denominator_note(stat.rows, f"assemblies assigned to {stat.name}"),
        }
        bioproject_dominance.append(dominance)
        if stat.rows >= 100 and dominance["top_5_bioproject_share_percent"] >= 30:
            bias_warnings.append(
                {
                    "scope": stat.name,
                    "scope_type": stat.rank,
                    "assemblies": stat.rows,
                    "bias_type": "BioProject dominance",
                    "severity": severity_from_share(dominance["top_5_bioproject_share_percent"]),
                    "metric_percent": dominance["top_5_bioproject_share_percent"],
                    "denominator": stat.rows,
                    "denominator_note": denominator_note(stat.rows, f"assemblies assigned to {stat.name}"),
                    "warning": (
                        f"{stat.name} has {dominance['top_5_bioproject_share_percent']}% of assemblies in the top five BioProjects. "
                        "Downstream analyses should consider BioProject-aware sampling."
                    ),
                }
            )
        for label, counter in (("country", stat.countries), ("host", stat.hosts), ("collection year", stat.years)):
            if not counter:
                continue
            top_label, top_count = counter.most_common(1)[0]
            share = percent(top_count, stat.rows)
            if share >= 50 and stat.rows >= 100:
                bias_warnings.append(
                    {
                        "scope": stat.name,
                        "scope_type": stat.rank,
                        "assemblies": stat.rows,
                        "bias_type": f"{label} dominance",
                        "severity": severity_from_share(share),
                        "metric_percent": share,
                        "denominator": stat.rows,
                        "denominator_note": denominator_note(stat.rows, f"assemblies assigned to {stat.name}"),
                        "warning": f"{stat.name} is dominated by {label} '{top_label}' ({share}% of assemblies).",
                    }
                )

    bioproject_dominance.sort(key=lambda row: (-float(row["top_5_bioproject_share_percent"]), -int(row["assemblies"])))
    bias_warnings.sort(key=lambda row: ({"severe": 0, "high": 1, "moderate": 2, "low": 3}.get(str(row["severity"]), 9), -int(row.get("assemblies") or 0), -float(row["metric_percent"])))

    correction_rows = [
        {
            "field": field,
            "primary_raw_value": raw,
            "raw_value": raw,
            "standardized_value": std,
            "evidence_source": evidence,
            "confidence": confidence,
            "correction_type": ctype,
            "records_rescued": count,
            "denominator": unique_total,
            "denominator_note": denominator_note(unique_total),
        }
        for (field, raw, std, evidence, confidence, ctype), count in correction_counter.most_common(100)
    ]

    field_confidence_summary = build_field_confidence_summary(field_confidence_counters, unique_total)
    validation_accuracy, validation_accuracy_available = load_validation_accuracy()
    case_studies = build_case_studies(taxon_stats, taxon_quality)
    top_bioprojects = []
    for rank, (project, count) in enumerate(bioproject_counter.most_common(100), start=1):
        dominant_taxon, dominant_taxon_count = bioproject_taxon_counter.get(project, Counter()).most_common(1)[0] if bioproject_taxon_counter.get(project) else ("not recorded", 0)
        top_bioprojects.append(
            {
                "rank": rank,
                "bioproject": project,
                "count": int(count),
                "percent": percent(count, unique_total),
                "dominant_taxon": dominant_taxon,
                "dominant_taxon_records": int(dominant_taxon_count),
                "dominant_taxon_share_within_bioproject_percent": percent(dominant_taxon_count, count),
                "denominator": unique_total,
                "denominator_note": denominator_note(unique_total),
            }
        )

    yearly_growth_rows = []
    cumulative = 0
    for year in sorted(year_counter):
        count = int(year_counter[year])
        cumulative += count
        yearly_growth_rows.append({"year": year, "assemblies": count, "cumulative_assemblies": cumulative, "denominator": unique_total, "denominator_note": denominator_note(unique_total)})

    priority_pathogens = load_priority_pathogens()
    taxon_quality_by_name = {normalized_key(row["taxon"]): row for row in taxon_quality}
    pathogen_rows = []
    for pathogen in priority_pathogens:
        name = pathogen.get("taxon_name", "")
        quality = taxon_quality_by_name.get(normalized_key(name))
        if quality:
            pathogen_rows.append({**pathogen, **quality})
        else:
            pathogen_rows.append({**pathogen, "assemblies": 0, "metadata_quality_score": "", "metadata_quality_grade": "Not available"})

    species_label_classes = taxonomy_label_summary_rows(
        species_label_class_assembly_counter,
        species_label_class_names,
        len(species_counter),
        unique_total,
    )
    catalog_species_label_classes = taxonomy_label_summary_rows(
        taxon_label_class_counter,
        taxon_label_class_names,
        sum(len(values) for values in taxon_label_class_names.values()),
    )

    app_commit_available = bool(app_commit and app_commit != "unknown")
    rule_manifest = standardization_rule_manifest()
    host_provenance = latest_host_standardization_provenance()
    geography_date_provenance = latest_geography_collection_date_provenance()
    tool_versions = tool_version_manifest()
    production_gate = rule_manifest.get("production_readiness_gate") or {}
    production_metrics = production_gate.get("metrics") or {}
    standardization_validation_checks = [
        ("file_errors", "file errors", 0),
        ("non_country_values_in_country_rows", "invalid standardized country rows", 0),
        ("country_continent_mismatch_rows", "country-continent mismatches", 0),
        ("country_subcontinent_mismatch_rows", "country-subcontinent mismatches", 0),
        ("invalid_sample_type_host_term_rows", "host-like sample-type rows", 0),
        ("noisy_isolation_source_broad_rows", "unapproved broad source rows", 0),
        ("body_site_leakage_values", "body-site leakage values", 0),
        ("disease_leakage_values", "disease/source leakage values", 0),
        ("raw_code_leakage_values", "raw code/text leakage values", 0),
        ("controlled_category_duplicate_keys", "controlled-category duplicate keys", 0),
        ("controlled_category_conflict_keys", "controlled-category conflict keys", 0),
        ("regression_tests_failed", "regression tests failed", 0),
    ]
    qa_checks = [
        {"check": "real_snapshot", "status": "pass", "detail": "Snapshot was generated from live metadata files, not demo values."},
        {"check": "app_commit_available", "status": "pass" if app_commit_available else "warning", "detail": app_commit if app_commit_available else "FETCHM_WEBAPP_GIT_COMMIT was not set during build."},
        {"check": "standardization_rules_fingerprinted", "status": "pass" if rule_manifest.get("version") != "not available" else "warning", "detail": f"{rule_manifest.get('version')} across {rule_manifest.get('total_rule_rows', 0):,} rule rows."},
        {"check": "standardization_audit_available", "status": "pass" if production_gate else "warning", "detail": f"Latest audit: {rule_manifest.get('latest_audit_timestamp', 'not available')}"},
        {"check": "standardization_production_ready", "status": "pass" if production_gate.get("production_ready") is True else "warning", "detail": "Production-readiness gate passed." if production_gate.get("production_ready") is True else "No passing production-readiness gate recorded."},
        {"check": "tool_versions_recorded", "status": "pass" if tool_versions.get("ncbi_datasets_version") != "not available" and tool_versions.get("taxonkit_version") != "not available" else "warning", "detail": f"NCBI Datasets: {tool_versions.get('ncbi_datasets_version')}; TaxonKit: {tool_versions.get('taxonkit_version')}"},
        {"check": "missing_values_excluded_from_biological_dominance", "status": "pass", "detail": "Missing tokens including absent/unknown/not applicable are excluded from host, country, source, and correction counters."},
        {"check": "correction_evidence_available", "status": "pass" if any(row.get("evidence_source") != "not recorded" for row in correction_rows[:30]) else "warning", "detail": "Top correction rows include evidence_source and confidence fields."},
        {"check": "metadata_files_scanned", "status": "pass" if files_scanned > 0 else "fail", "detail": f"{files_scanned} files scanned; {files_skipped} files skipped."},
        {"check": "duplicate_rows_recorded", "status": "pass", "detail": f"{duplicate_rows} duplicate rows skipped by Assembly Accession."},
        {"check": "manual_validation_accuracy", "status": "pass" if validation_accuracy_available else "warning", "detail": "Manual validation accuracy CSV was summarized." if validation_accuracy_available else "No manual validation_records.csv found; validation accuracy panel reports the expected input format."},
    ]
    for metric_key, label, expected in standardization_validation_checks:
        raw_value = production_metrics.get(metric_key)
        try:
            numeric_value = int(float(str(raw_value)))
        except (TypeError, ValueError):
            numeric_value = None
        qa_checks.append(
            {
                "check": metric_key,
                "status": "pass" if numeric_value == expected else "warning",
                "detail": f"{label}: {raw_value if raw_value not in (None, '') else 'not recorded'}",
            }
        )
    qa_checks.append(
        {
            "check": "review_needed_counts",
            "status": "pass",
            "detail": (
                f"Host review-needed rows: {production_metrics.get('host_review_needed_rows', 'not recorded')}; "
                f"source-like mapped host rows for review: {production_metrics.get('source_like_mapped_host_rows_for_review', 'not recorded')}; "
                f"source-like unmapped host rows for review: {production_metrics.get('source_like_unmapped_host_rows_for_review', 'not recorded')}"
            ),
        }
    )
    qa_status = "fail" if any(row["status"] == "fail" for row in qa_checks) else "warning" if any(row["status"] == "warning" for row in qa_checks) else "pass"

    summary = {
        "snapshot_id": snapshot_id,
        "generated_at": utc_now(),
        "is_demo": False,
        "overview": {
            "unique_assemblies": unique_total,
            "metadata_rows_scanned": metadata_rows_scanned,
            "duplicate_rows_skipped": duplicate_rows,
            "metadata_files_scanned": files_scanned,
            "metadata_files_skipped": files_skipped,
            "genera_observed": len(genus_counter),
            "species_observed": len(species_counter),
            "canonical_species_observed": len(species_label_class_names.get("canonical_species", set())),
            "unresolved_species_level_labels_observed": len(species_label_class_names.get("unresolved_species_level_label", set())),
            "provisional_species_labels_observed": len(species_label_class_names.get("provisional_taxonomic_label", set())),
            "noncanonical_species_labels_observed": len(species_label_class_names.get("noncanonical_species_label", set())),
            "species_pages_observed": sum(len(values) for values in taxon_label_class_names.values()),
            "canonical_species_pages": len(taxon_label_class_names.get("canonical_species", set())),
            "unresolved_species_level_label_pages": len(taxon_label_class_names.get("unresolved_species_level_label", set())),
            "provisional_taxonomic_label_pages": len(taxon_label_class_names.get("provisional_taxonomic_label", set())),
            "noncanonical_species_label_pages": len(taxon_label_class_names.get("noncanonical_species_label", set())),
            "countries_observed": len(country_counter),
            "hosts_observed": len(host_counter),
            "bioprojects_observed": len(bioproject_counter),
            "qc_ready_taxa": len(qc_ready_taxa),
        },
        "methods": {
            "app_version": app_version,
            "app_commit": app_commit,
            "generation_scope": (
                f"Activated canonical bacterial inventory snapshot {source_snapshot_id} was scanned once at assembly level after standardization and reconciliation."
                if canonical_root_source
                else "Ready FetchM taxa with metadata_status='ready' and a metadata_clean_path were scanned after metadata update and standardization refresh completion."
            ),
            "source_snapshot_id": source_snapshot_id or "legacy-managed-taxa",
            "canonical_root_source": bool(canonical_root_source),
            "duplicate_rule": (
                "Unique assemblies are counted by Assembly Accession from the activated canonical root inventory; each assembly is scanned once."
                if canonical_root_source
                else "Unique assemblies are counted by Assembly Accession. Species-level metadata rows are preferred over genus-level rows; newer synced taxa are scanned first within each rank."
            ),
            "field_mappings": "Country: Geographic Location/Country_Raw -> Country; Host: Host -> Host_SD; Isolation source: Isolation Source -> Isolation_Source_SD; sample type and environment raw comparisons use raw sample/environment fields when present, while standardized coverage is counted from Sample_Type_SD and Environment_*_SD; collection year from collection-date metadata: Collection Date standardized to a 4-digit year, with Collection_Date_Source/Evidence/Recovery_Status retained when available; growth year: Assembly Release Date.",
            "missing_value_rule": "Empty, unknown, not collected, not applicable, unidentified, and similarly non-informative values are treated as unusable metadata.",
            "taxonomy_label_policy": "Species-level labels are retained comprehensively for search and metadata access, but classified as canonical species, unresolved species-level labels, provisional taxonomic labels, or non-canonical species labels for honest reporting.",
            "qc_ready_rule": "At least 100 assemblies, >=70% standardized country completeness, >=70% host/source completeness, and >=50% collection-year completeness.",
            "bias_score_formulas": "Dominance scores are calculated as the top category share among assemblies in scope: top 1/5/10 BioProject share, top country share, top host share, and top collection-year share. Warning severity uses low <25%, moderate 25-49.99%, high 50-74.99%, and severe >=75%.",
            "metadata_quality_score": "0.20 country + 0.20 host/source + 0.15 collection year + 0.15 assembly quality + 0.10 BioProject diversity + 0.10 isolation source + 0.10 standardization confidence.",
            "retrieval_date": latest_taxon_synced_at or iso_from_timestamp(latest_metadata_mtime),
            "biosample_enrichment_date": latest_taxon_synced_at or iso_from_timestamp(latest_metadata_mtime),
            "standardization_refresh_date": iso_from_timestamp(latest_metadata_mtime),
            "files_scanned": files_scanned,
            "files_skipped": files_skipped,
            "metadata_rows_scanned": metadata_rows_scanned,
            "duplicate_rows_skipped": duplicate_rows,
            "final_unique_assemblies": unique_total,
            "assembly_source": "all public assemblies represented by the current FetchM metadata snapshot; exact original GenBank/RefSeq/all request is not stored per legacy row",
            "standardization_rule_version": rule_manifest.get("version", "not available"),
            "standardization_rule_rows": rule_manifest.get("total_rule_rows", 0),
            "standardization_approved_rule_rows": rule_manifest.get("approved_rule_rows", 0),
            "standardization_review_or_nonapproved_rule_rows": rule_manifest.get("review_or_nonapproved_rule_rows", 0),
            "standardization_rule_files": rule_manifest.get("files", []),
            "host_standardization_provenance": host_provenance,
            "geography_collection_date_provenance": geography_date_provenance,
            "latest_standardization_audit_timestamp": rule_manifest.get("latest_audit_timestamp", "not available"),
            "latest_standardization_audit_git_commit": rule_manifest.get("latest_audit_git_commit", "not available"),
            "latest_standardization_audit_code_version": rule_manifest.get("latest_audit_code_version", "not available"),
            "ncbi_datasets_version": tool_versions.get("ncbi_datasets_version", "not available"),
            "taxonkit_version": tool_versions.get("taxonkit_version", "not available"),
            "caution": "Global Insights describe representation within public genome repositories, not true global bacterial abundance, disease burden, or environmental prevalence.",
        },
        "qa": {
            "status": qa_status,
            "checks": qa_checks,
            "manuscript_readiness": {
                "real_snapshot": True,
                "app_commit_available": app_commit_available,
                "qa_gate_passed": qa_status == "pass",
                "rule_version_available": rule_manifest.get("version") != "not available",
                "validation_summary_available": bool(production_gate) or validation_accuracy_available,
                "tool_versions_available": tool_versions.get("ncbi_datasets_version") != "not available" and tool_versions.get("taxonkit_version") != "not available",
                "all_figures_exported": False,
                "source_data_exported": True,
                "missing_values_not_biological_categories": True,
                "snapshot_archived": True,
            "checksums_available": False,
            "publication_export_audit_passed": False,
            },
        },
        "taxonomic_landscape": {
            "top_genera": top_rows(genus_counter, unique_total, 25),
            "top_species": top_rows(species_counter, unique_total, 25),
            "species_label_classes": species_label_classes,
            "catalog_species_label_classes": catalog_species_label_classes,
            "top_10_genus_share_percent": percent(sum(count for _, count in genus_counter.most_common(10)), unique_total),
        },
        "geographic_bias": {
            "countries": top_rows(country_counter, unique_total, 50),
            "continents": top_rows(continent_counter, unique_total, 20),
            "subcontinents": top_rows(subcontinent_counter, unique_total, 30),
            "non_country_geography": top_rows(non_country_geography_counter, unique_total, 20),
            "income_groups": [],
            "income_group_note": "World Bank income-group representation is not calculated until a versioned country-to-income lookup is configured.",
        },
        "host_source_bias": {
            "hosts": top_rows(host_counter, unique_total, 50),
            "sources": top_rows(source_counter, unique_total, 50),
            "host_categories": top_rows(host_category_counter, unique_total, 20),
        },
        "assembly_quality": {
            "assembly_levels": top_rows(assembly_level_counter, unique_total, 20),
        },
        "metadata_completeness": metadata_completeness,
        "standardization_impact": {
            "top_corrections": correction_rows[:30],
            "confidence_methods": top_rows(confidence_counter, sum(confidence_counter.values()), 30),
            "mapped_records": int(sum(row["records_rescued"] for row in correction_rows)),
        },
        "yearly_growth": yearly_growth_rows,
        "metadata_quality": taxon_quality[:100],
        "qc_ready_taxa": qc_ready_taxa[:100],
        "metadata_readiness_tiers": readiness_tiers,
        "field_confidence_summary": field_confidence_summary,
        "validation_accuracy": validation_accuracy,
        "case_studies": case_studies,
        "top_bioprojects": top_bioprojects[:100],
        "bioproject_dominance": bioproject_dominance[:100],
        "bias_warnings": bias_warnings[:100],
        "pathogen_insights": pathogen_rows,
        "downloads": {
            "summary_json": "summary.json",
            "simulator_records": "tables/simulator_records.csv",
            "top_genera": "tables/top_genera.csv",
            "top_species": "tables/top_species.csv",
            "countries": "tables/countries.csv",
            "continents": "tables/continents.csv",
            "subcontinents": "tables/subcontinents.csv",
            "non_country_geography": "tables/non_country_geography.csv",
            "host_categories": "tables/host_categories.csv",
            "metadata_completeness": "tables/metadata_completeness.csv",
            "metadata_quality": "tables/metadata_quality.csv",
            "metadata_readiness_tiers": "tables/metadata_readiness_tiers.csv",
            "field_confidence_summary": "tables/field_confidence_summary.csv",
            "validation_accuracy": "tables/validation_accuracy.csv",
            "case_studies": "tables/case_studies.csv",
            "top_bioprojects": "tables/top_bioprojects.csv",
            "species_label_classes": "tables/species_label_classes.csv",
            "catalog_species_label_classes": "tables/catalog_species_label_classes.csv",
            "bias_warnings": "tables/bias_warnings.csv",
            "top_corrections": "tables/top_corrections.csv",
            "yearly_growth": "tables/yearly_growth.csv",
            "qc_ready_taxa": "tables/qc_ready_taxa.csv",
            "bioproject_dominance": "tables/bioproject_dominance.csv",
            "pathogen_insights": "tables/pathogen_insights.csv",
        },
    }
    summary["manuscript"] = build_narrative(summary)

    top_row_fields = ["rank", "label", "count", "percent", "denominator", "denominator_note"]
    write_csv(table_dir / "top_genera.csv", summary["taxonomic_landscape"]["top_genera"], top_row_fields)
    write_csv(table_dir / "top_species.csv", summary["taxonomic_landscape"]["top_species"], top_row_fields)
    label_class_fields = ["label_class", "label", "description", "labels", "label_percent", "assemblies", "assembly_percent", "denominator", "denominator_note", "assembly_denominator", "assembly_denominator_note"]
    write_csv(table_dir / "species_label_classes.csv", summary["taxonomic_landscape"].get("species_label_classes", []), label_class_fields)
    write_csv(table_dir / "catalog_species_label_classes.csv", summary["taxonomic_landscape"].get("catalog_species_label_classes", []), label_class_fields)
    write_csv(table_dir / "countries.csv", summary["geographic_bias"]["countries"], top_row_fields)
    write_csv(table_dir / "continents.csv", summary["geographic_bias"]["continents"], top_row_fields)
    write_csv(table_dir / "subcontinents.csv", summary["geographic_bias"]["subcontinents"], top_row_fields)
    write_csv(table_dir / "non_country_geography.csv", summary["geographic_bias"].get("non_country_geography", []), top_row_fields)
    write_csv(table_dir / "host_categories.csv", summary["host_source_bias"]["host_categories"], top_row_fields)
    write_csv(
        table_dir / "metadata_completeness.csv",
        metadata_completeness,
        [
            "field",
            "raw_usable",
            "raw_usable_percent",
            "standardized_usable",
            "standardized_usable_percent",
            "both_raw_and_standardized_usable",
            "standardized_only_rescued",
            "raw_only_unresolved",
            "neither_usable",
            "changed_mapping_count",
            "standardized_gain_percentage_points",
            "denominator",
            "denominator_note",
        ],
    )
    write_csv(table_dir / "metadata_quality.csv", taxon_quality, list(taxon_quality[0].keys()) if taxon_quality else ["taxon", "rank", "assemblies"])
    write_csv(table_dir / "metadata_readiness_tiers.csv", readiness_tiers, ["tier", "taxa", "definition", "denominator", "denominator_note"])
    write_csv(table_dir / "field_confidence_summary.csv", field_confidence_summary, ["field", "confidence_status", "count", "percent", "denominator", "denominator_note"])
    write_csv(table_dir / "validation_accuracy.csv", validation_accuracy, ["field", "validation_records", "precision_percent", "false_positive_rate_percent", "unresolved_rate_percent", "common_error_types", "validation_source", "status"])
    write_csv(table_dir / "case_studies.csv", case_studies, sorted({key for row in case_studies for key in row.keys()}) if case_studies else ["taxon", "status", "assemblies"])
    write_csv(table_dir / "top_bioprojects.csv", top_bioprojects, ["rank", "bioproject", "count", "percent", "dominant_taxon", "dominant_taxon_records", "dominant_taxon_share_within_bioproject_percent", "denominator", "denominator_note"])
    write_csv(table_dir / "bias_warnings.csv", bias_warnings, ["scope", "scope_type", "assemblies", "bias_type", "severity", "metric_percent", "denominator", "denominator_note", "warning"])
    write_csv(table_dir / "top_corrections.csv", correction_rows, ["field", "primary_raw_value", "raw_value", "standardized_value", "evidence_source", "confidence", "correction_type", "records_rescued", "denominator", "denominator_note"])
    write_csv(table_dir / "yearly_growth.csv", yearly_growth_rows, ["year", "assemblies", "cumulative_assemblies", "denominator", "denominator_note"])
    write_csv(table_dir / "qc_ready_taxa.csv", qc_ready_taxa, list(qc_ready_taxa[0].keys()) if qc_ready_taxa else ["taxon", "rank", "assemblies"])
    write_csv(table_dir / "bioproject_dominance.csv", bioproject_dominance, ["taxon", "rank", "assemblies", "top_1_bioproject_share_percent", "top_5_bioproject_share_percent", "top_10_bioproject_share_percent", "top_bioproject", "denominator", "denominator_note"])
    write_csv(table_dir / "pathogen_insights.csv", pathogen_rows, sorted({key for row in pathogen_rows for key in row.keys()}) if pathogen_rows else ["taxon_name", "rank", "group", "notes"])

    generate_publication_exports(summary, tmp_dir, table_dir)
    summary.setdefault("downloads", {}).update(
        {
            "archive_manifest": "manifest.json",
            "provenance_manifest": "provenance/manifest.json",
            "software_versions": "provenance/software_versions.json",
            "rule_fingerprint": "provenance/rule_fingerprint.json",
            "qa_report": "provenance/qa_report.json",
            "publication_export_audit": "provenance/publication_export_audit.json",
            "checksums_sha256": "checksums.sha256",
        }
    )
    summary.setdefault("qa", {}).setdefault("manuscript_readiness", {})["checksums_available"] = True
    completed_at = utc_now()
    status_path.write_text(json.dumps({"snapshot_id": snapshot_id, "status": "completed", "completed_at": completed_at}, indent=2), encoding="utf-8")
    write_snapshot_manifest(summary, tmp_dir)
    (tmp_dir / "summary.json").write_text(json.dumps(summary, indent=2), encoding="utf-8")
    write_checksums(tmp_dir)
    if snapshot_dir.exists():
        shutil.rmtree(snapshot_dir)
    tmp_dir.rename(snapshot_dir)
    latest = {"snapshot_id": snapshot_id, "summary_path": str(snapshot_dir / "summary.json"), "generated_at": summary["generated_at"], "is_demo": False}
    (output_root / "latest.json").write_text(json.dumps(latest, indent=2), encoding="utf-8")
    return summary


def generate_demo_snapshot(output_root: Path, *, app_version: str, app_commit: str, snapshot_id: str | None = None) -> dict[str, Any]:
    snapshot_id = snapshot_id or "demo_global_insights"
    output_root.mkdir(parents=True, exist_ok=True)
    snapshot_dir = output_root / "snapshots" / safe_slug(snapshot_id)
    if snapshot_dir.exists():
        shutil.rmtree(snapshot_dir)
    table_dir = snapshot_dir / "tables"
    table_dir.mkdir(parents=True, exist_ok=True)
    app_commit_available = bool(app_commit and app_commit != "unknown")
    qa_checks = [
        {"check": "real_snapshot", "status": "warning", "detail": "Demo snapshot uses placeholder values and must not be cited."},
        {"check": "app_commit_available", "status": "pass" if app_commit_available else "warning", "detail": app_commit if app_commit_available else "FETCHM_WEBAPP_GIT_COMMIT was not set during build."},
        {"check": "missing_values_excluded_from_biological_dominance", "status": "pass", "detail": "Demo follows the same missing-value handling policy as live snapshots."},
        {"check": "publication_exports_generated", "status": "pending", "detail": "Publication exports are generated after demo summary assembly."},
    ]
    qa_status = "warning"

    summary = {
        "snapshot_id": snapshot_id,
        "generated_at": utc_now(),
        "is_demo": True,
        "overview": {
            "unique_assemblies": 2846213,
            "metadata_rows_scanned": 3120448,
            "duplicate_rows_skipped": 274235,
            "metadata_files_scanned": 5066,
            "metadata_files_skipped": 0,
            "genera_observed": 5814,
            "species_observed": 48392,
            "canonical_species_observed": 42110,
            "unresolved_species_level_labels_observed": 5142,
            "provisional_species_labels_observed": 820,
            "noncanonical_species_labels_observed": 320,
            "species_pages_observed": 48392,
            "canonical_species_pages": 42110,
            "unresolved_species_level_label_pages": 5142,
            "provisional_taxonomic_label_pages": 820,
            "noncanonical_species_label_pages": 320,
            "countries_observed": 184,
            "hosts_observed": 12140,
            "bioprojects_observed": 152706,
            "qc_ready_taxa": 1286,
        },
        "methods": {
            "app_version": app_version,
            "app_commit": app_commit,
            "generation_scope": "DEMO: ready standardized metadata files after refresh completion.",
            "duplicate_rule": "DEMO: unique assemblies are counted by Assembly Accession.",
            "field_mappings": "DEMO: raw country/host/source fields are compared with FetchM standardized country/host/source fields.",
            "missing_value_rule": "DEMO: non-informative metadata values are treated as unusable.",
            "taxonomy_label_policy": "DEMO: species-level labels are retained for search but classified separately from canonical species.",
            "qc_ready_rule": "DEMO: >=100 assemblies and adequate standardized metadata completeness.",
            "bias_score_formulas": "DEMO: dominance shares use top BioProject, country, host, and year representation.",
            "metadata_quality_score": "DEMO: weighted country, host/source, year, quality, BioProject diversity, source, and confidence components.",
            "caution": "DEMO DATA - NOT REAL RESULTS. Repository representation is not global biological prevalence.",
        },
        "taxonomic_landscape": {
            "top_genera": [
                {"rank": 1, "label": "Escherichia", "count": 318420, "percent": 11.19},
                {"rank": 2, "label": "Salmonella", "count": 284611, "percent": 10.0},
                {"rank": 3, "label": "Staphylococcus", "count": 201984, "percent": 7.1},
                {"rank": 4, "label": "Klebsiella", "count": 176392, "percent": 6.2},
                {"rank": 5, "label": "Mycobacterium", "count": 144210, "percent": 5.07},
            ],
            "top_species": [
                {"rank": 1, "label": "Escherichia coli", "count": 287441, "percent": 10.1},
                {"rank": 2, "label": "Salmonella enterica", "count": 241882, "percent": 8.5},
                {"rank": 3, "label": "Staphylococcus aureus", "count": 184901, "percent": 6.5},
            ],
            "species_label_classes": [
                {"label_class": "canonical_species", "label": "Canonical species", "description": TAXONOMY_LABEL_CLASS_DESCRIPTIONS["canonical_species"], "labels": 42110, "label_percent": 87.02, "assemblies": 2600000, "assembly_percent": 91.35},
                {"label_class": "unresolved_species_level_label", "label": "Unresolved species-level label", "description": TAXONOMY_LABEL_CLASS_DESCRIPTIONS["unresolved_species_level_label"], "labels": 5142, "label_percent": 10.63, "assemblies": 210000, "assembly_percent": 7.38},
            ],
            "catalog_species_label_classes": [
                {"label_class": "canonical_species", "label": "Canonical species", "description": TAXONOMY_LABEL_CLASS_DESCRIPTIONS["canonical_species"], "labels": 42110, "label_percent": 87.02},
                {"label_class": "unresolved_species_level_label", "label": "Unresolved species-level label", "description": TAXONOMY_LABEL_CLASS_DESCRIPTIONS["unresolved_species_level_label"], "labels": 5142, "label_percent": 10.63},
            ],
            "top_10_genus_share_percent": 49.6,
        },
        "geographic_bias": {
            "countries": [{"rank": 1, "label": "United States", "count": 702112, "percent": 24.67}],
            "continents": [
                {"rank": 1, "label": "North America", "count": 948210, "percent": 33.31},
                {"rank": 2, "label": "Europe", "count": 821403, "percent": 28.86},
                {"rank": 3, "label": "Asia", "count": 612992, "percent": 21.54},
            ],
            "subcontinents": [],
        },
        "host_source_bias": {
            "hosts": [{"rank": 1, "label": "Human", "count": 1042770, "percent": 36.64}],
            "sources": [{"rank": 1, "label": "feces/stool", "count": 381119, "percent": 13.39}],
            "host_categories": [
                {"rank": 1, "label": "Human-associated", "count": 1042770, "percent": 36.64},
                {"rank": 2, "label": "Animal-associated", "count": 428119, "percent": 15.04},
            ],
        },
        "assembly_quality": {"assembly_levels": [{"rank": 1, "label": "Contig", "count": 1800000, "percent": 63.24}]},
        "metadata_completeness": [
            {"field": "Country", "raw_usable": 1702114, "raw_usable_percent": 59.8, "standardized_usable": 2436890, "standardized_usable_percent": 85.6, "rescued_records": 734776, "gain_percentage_points": 25.8},
            {"field": "Host", "raw_usable": 1371947, "raw_usable_percent": 48.2, "standardized_usable": 2189338, "standardized_usable_percent": 76.9, "rescued_records": 817391, "gain_percentage_points": 28.7},
            {"field": "Isolation source", "raw_usable": 1238090, "raw_usable_percent": 43.5, "standardized_usable": 2026530, "standardized_usable_percent": 71.2, "rescued_records": 788440, "gain_percentage_points": 27.7},
        ],
        "standardization_impact": {
            "top_corrections": [
                {"field": "Country", "raw_value": "USA", "standardized_value": "United States", "records_rescued": 82441},
                {"field": "Host", "raw_value": "Homo sapiens", "standardized_value": "Human", "records_rescued": 64220},
            ],
            "confidence_methods": [],
        },
        "yearly_growth": [
            {"year": "2022", "assemblies": 310000, "cumulative_assemblies": 2010000},
            {"year": "2023", "assemblies": 386000, "cumulative_assemblies": 2396000},
            {"year": "2024", "assemblies": 420000, "cumulative_assemblies": 2816000},
        ],
        "metadata_quality": [
            {"taxon": "Salmonella enterica", "rank": "species", "assemblies": 241882, "metadata_quality_score": 91.2, "metadata_quality_grade": "Excellent"},
            {"taxon": "Escherichia coli", "rank": "species", "assemblies": 287441, "metadata_quality_score": 83.4, "metadata_quality_grade": "Good"},
        ],
        "qc_ready_taxa": [],
        "bioproject_dominance": [],
        "bias_warnings": [
            {"scope": "Salmonella enterica", "scope_type": "species", "bias_type": "BioProject dominance", "severity": "moderate", "metric_percent": 31.4, "warning": "DEMO: top five BioProjects account for 31.4% of assemblies."}
        ],
        "pathogen_insights": [],
        "downloads": {},
        "qa": {
            "status": qa_status,
            "checks": qa_checks,
            "manuscript_readiness": {
                "real_snapshot": False,
                "app_commit_available": app_commit_available,
                "qa_gate_passed": False,
                "all_figures_exported": False,
                "source_data_exported": False,
                "missing_values_not_biological_categories": True,
                "snapshot_archived": True,
            "checksums_available": False,
            "publication_export_audit_passed": False,
            },
        },
    }
    summary["manuscript"] = build_narrative(summary)
    generate_publication_exports(summary, snapshot_dir, table_dir)
    (snapshot_dir / "summary.json").write_text(json.dumps(summary, indent=2), encoding="utf-8")
    (snapshot_dir / "status.json").write_text(json.dumps({"snapshot_id": snapshot_id, "status": "completed", "completed_at": utc_now()}, indent=2), encoding="utf-8")
    (output_root / "latest.json").write_text(json.dumps({"snapshot_id": snapshot_id, "summary_path": str(snapshot_dir / "summary.json"), "generated_at": summary["generated_at"], "is_demo": True}, indent=2), encoding="utf-8")
    return summary


def matches_filter(value: str, expected: str) -> bool:
    if not expected:
        return True
    return expected.casefold() in normalize(value).casefold()


def run_standardization_simulator(records_path: Path, filters: dict[str, Any], *, limit_examples: int = 25) -> dict[str, Any]:
    if not records_path.exists():
        return {"available": False, "error": "Simulator records are not available for this snapshot."}
    raw_count = 0
    standardized_count = 0
    rescued_count = 0
    examples: list[dict[str, str]] = []
    country = normalize(filters.get("country"))
    host = normalize(filters.get("host"))
    taxon = normalize(filters.get("taxon"))
    assembly_level = normalize(filters.get("assembly_level"))
    year_from = parse_year(str(filters.get("year_from") or ""))
    year_to = parse_year(str(filters.get("year_to") or ""))
    with records_path.open(newline="", encoding="utf-8", errors="replace") as handle:
        for row in csv.DictReader(handle):
            year = row.get("collection_year", "")
            if year_from and (not year or year < year_from):
                continue
            if year_to and (not year or year > year_to):
                continue
            if taxon and not (matches_filter(row.get("taxon", ""), taxon) or matches_filter(row.get("species", ""), taxon) or matches_filter(row.get("genus", ""), taxon)):
                continue
            if assembly_level and not matches_filter(row.get("assembly_level", ""), assembly_level):
                continue
            raw_match = matches_filter(row.get("raw_country", ""), country) and matches_filter(row.get("raw_host", ""), host)
            std_match = matches_filter(row.get("standardized_country", ""), country) and (
                matches_filter(row.get("standardized_host", ""), host) or matches_filter(row.get("standardized_source", ""), host)
            )
            if raw_match:
                raw_count += 1
            if std_match:
                standardized_count += 1
            if std_match and not raw_match:
                rescued_count += 1
                if len(examples) < limit_examples:
                    examples.append(
                        {
                            "assembly_accession": row.get("assembly_accession", ""),
                            "taxon": row.get("taxon", ""),
                            "raw_country": row.get("raw_country", ""),
                            "standardized_country": row.get("standardized_country", ""),
                            "raw_host": row.get("raw_host", ""),
                            "standardized_host": row.get("standardized_host", ""),
                        }
                    )
    return {
        "available": True,
        "raw_count": raw_count,
        "standardized_count": standardized_count,
        "rescued_count": rescued_count,
        "gain_percent": percent(rescued_count, raw_count) if raw_count else 0.0,
        "examples": examples,
    }
